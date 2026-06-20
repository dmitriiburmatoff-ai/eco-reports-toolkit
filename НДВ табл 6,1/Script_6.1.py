import os
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# ==============================================================================
# 1. Вспомогательные функции для парсинга и форматирования
# ==============================================================================

def parse_float(val):
    """Пытается преобразовать строку в float. Игнорирует прочерки и пустоту."""
    if val is None:
        return None
    val_str = str(val).strip().replace('\xa0', '').replace(' ', '').replace(',', '.')
    if val_str in ('', '-', '–', '—'):
        return None
    try:
        return float(val_str)
    except ValueError:
        return None

def format_value(num):
    """Форматирует число: >= 0.01: два знака, < 0.01: экспоненциальная форма."""
    if num is None:
        return '---'
    if num >= 0.01:
        return f"{num:.2f}".replace('.', ',')
    else:
        return f"{num:.2E}".replace('.', ',')

def detect_doc_type(file_path):
    """Определяет тип документа по тексту в первой таблице, а не по имени файла."""
    try:
        doc = docx.Document(file_path)
        if not doc.tables:
            return "UNKNOWN"
        # Собираем весь текст из первой таблицы в одну строку
        text = " ".join(cell.text for row in doc.tables[0].rows for cell in row.cells).lower()
        if "расчет рассеивания" in text:
            return "MR"
        if "среднесуточных" in text:
            return "SS"
        return "UNKNOWN"
    except Exception:
        return "UNKNOWN"

# ==============================================================================
# 2. Функция извлечения данных
# ==============================================================================

def extract_max_t4_data(file_path):
    """
    Извлекает данные для всех веществ, находя максимальную концентрацию (д. ПДК, индекс 4) 
    только для строк, где Тип точки (последняя ячейка) равен '4'.
    """
    try:
        doc = docx.Document(file_path)
        results = {}
        current_block = None
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                # Начало нового блока вещества
                if 'Вещество:' in ''.join(cells):
                    if current_block:
                        results[current_block['code']] = current_block

                    raw = cells[0]
                    after_colon = raw.split(':', 1)[1].strip()
                    parts = after_colon.split('\n', 1)
                    code = parts[0].strip()
                    name = parts[1].strip() if len(parts) > 1 else ''

                    if not code:
                        current_block = None
                        continue

                    current_block = {
                        'code': code,
                        'name': name,
                        'type4_rows': []
                    }
                    continue

                if current_block is None:
                    continue

                # Пропускаем строки-заголовки
                if cells[0] in ('', '№') or 'Тип' in cells[-1]:
                    continue

                # Строки данных: тип точки в последней ячейке
                point_type = cells[-1].strip()
                if point_type == '4':
                    point_num = cells[0]
                    # ИСПРАВЛЕНИЕ: берем индекс 4 (Концентр. д. ПДК)
                    conc_str = cells[4] if len(cells) > 4 else ''
                    
                    current_block['type4_rows'].append({
                        'номер': point_num,
                        'конц_str': conc_str
                    })

        # Не забываем последний блок
        if current_block:
            results[current_block['code']] = current_block

        # Формируем итоговые записи с флагом валидности данных
        final_data = {}
        for code, block in results.items():
            best_num = ''
            best_val = None
            has_valid = False

            for r in block['type4_rows']:
                val = parse_float(r['конц_str'])
                if val is not None:
                    has_valid = True
                    if best_val is None or val > best_val:
                        best_val = val
                        best_num = r['номер']

            final_data[code] = {
                'name': block['name'],
                'point': best_num,
                'conc': best_val,
                'has_valid': has_valid
            }

        return final_data

    except Exception as e:
        print(f"❌ Ошибка при обработке файла {file_path}: {e}")
        return {}

# ==============================================================================
# 3. Функции для форматирования Word-документа
# ==============================================================================

def set_cell_text(cell, text, font_name='Times New Roman', font_size=10,
                  bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def save_to_docx(data, output_path):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(20)
        section.right_margin = Pt(15)

    if not data:
        doc.add_paragraph("Нет данных для вывода.")
        doc.save(output_path)
        return

    total_rows = 3 + 1 + len(data)
    table = doc.add_table(rows=total_rows, cols=6)
    table.style = 'Table Grid'

    # Заголовок, строка 0
    table.cell(0, 0).merge(table.cell(2, 0))
    set_cell_text(table.cell(0, 0), 'Загрязняющее вещество, код и наименование', bold=True)
    table.cell(0, 1).merge(table.cell(2, 1))
    set_cell_text(table.cell(0, 1), '№ РТ', bold=True)
    table.cell(0, 2).merge(table.cell(0, 5))
    set_cell_text(table.cell(0, 2), 'Расчетная максимальная приземная концентрация, в долях ПДК', bold=True)

    # Заголовок, строка 1
    table.cell(1, 2).merge(table.cell(2, 2))
    set_cell_text(table.cell(1, 2), 'Концентрация (доля ПДК)', bold=True)
    table.cell(1, 3).merge(table.cell(2, 3))
    set_cell_text(table.cell(1, 3), 'с увеличением на 20%', bold=True)
    table.cell(1, 4).merge(table.cell(2, 4))
    set_cell_text(table.cell(1, 4), 'с увеличением на 40%', bold=True)
    table.cell(1, 5).merge(table.cell(2, 5))
    set_cell_text(table.cell(1, 5), 'с увеличением на 60%', bold=True)

    # Нумерация, строка 3
    for ci, num in enumerate(['1', '2', '3', '4', '5', '6']):
        set_cell_text(table.cell(3, ci), num)

    # Данные
    for i, record in enumerate(data):
        row = table.rows[4 + i]
        set_cell_text(row.cells[0], record['Вещество'], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[1], str(record['НомерТочки']))
        set_cell_text(row.cells[2], format_value(record['Конц']))
        set_cell_text(row.cells[3], format_value(record['Конц_120']))
        set_cell_text(row.cells[4], format_value(record['Конц_140']))
        set_cell_text(row.cells[5], format_value(record['Конц_160']))

    doc.save(output_path)

# ==============================================================================
# 4. Основная логика
# ==============================================================================

def main():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 1. Поиск файлов (игнорируем временные, выходные и старые НДВ)
    docx_files = [f for f in os.listdir(current_dir) 
                  if f.endswith('.docx') 
                  and not f.startswith('~') 
                  and not f.startswith('output_') 
                  and not f.startswith('НДВ_')]
    
    print(f"🔍 Поиск файлов в папке: {current_dir}")
    print(f"   Найдено кандидатов: {len(docx_files)} -> {docx_files}")
    
    if len(docx_files) == 0:
        print("❌ Не найдено ни одного подходящего файла .docx в папке.")
        return

    # 2. Определение типов файлов по СОДЕРЖИМОМУ
    mr_file = None
    ss_file = None
    
    for fname in docx_files:
        fpath = os.path.join(current_dir, fname)
        dtype = detect_doc_type(fpath)
        if dtype == "MR":
            if mr_file is None:  # Берем первый найденный
                mr_file = fpath
                print(f"✅ Найден файл МР (по содержимому): {fname}")
            else:
                print(f"⚠️ Найден второй файл МР, игнорируем: {fname}")
        elif dtype == "SS":
            if ss_file is None:  # Берем первый найденный
                ss_file = fpath
                print(f"✅ Найден файл СС (по содержимому): {fname}")
            else:
                print(f"⚠️ Найден второй файл СС, игнорируем: {fname}")
        else:
            print(f"⚠️ Не удалось определить тип файла (пропущен): {fname}")

    if not mr_file:
        print("❌ Не удалось найти файл типа МР (Расчет рассеивания). Без него обработка невозможна.")
        return

    if not ss_file:
        print("⚠️ Не удалось найти файл типа СС. Резервное копирование данных будет невозможно, будут использованы только данные из МР.")

    # 3. Извлечение данных
    print("\n🔄 Обработка файла МР...")
    mr_data = extract_max_t4_data(mr_file)
    print(f"   Найдено веществ в МР: {len(mr_data)}")

    ss_data = {}
    if ss_file:
        print("🔄 Обработка файла СС...")
        ss_data = extract_max_t4_data(ss_file)
        print(f"   Найдено веществ в СС: {len(ss_data)}")

    # 4. Слияние данных по правилу: если в МР прочерк, берем из СС
    merged_data = []
    fallback_count = 0

    for code, mr_info in mr_data.items():
        if not mr_info['has_valid']:
            # В МР нет валидных данных (прочерки), пробуем взять из СС
            if ss_file and code in ss_data and ss_data[code]['has_valid']:
                final_conc = ss_data[code]['conc']
                final_point = ss_data[code]['point']
                name = ss_data[code]['name']
                fallback_count += 1
            else:
                # Нет данных и в СС, пропускаем вещество
                continue
        else:
            # В МР есть валидные данные, используем их
            final_conc = mr_info['conc']
            final_point = mr_info['point']
            name = mr_info['name']

        merged_data.append({
            'Вещество': f"{code} {name}" if name else code,
            'НомерТочки': final_point,
            'Конц': final_conc,
            'Конц_120': final_conc * 1.20,
            'Конц_140': final_conc * 1.40,
            'Конц_160': final_conc * 1.60,
        })

    # Сортировка по коду вещества для аккуратного вида
    merged_data.sort(key=lambda x: x['Вещество'])

    # 5. Сохранение результата
    output_filename = 'НДВ_6.1.docx'
    output_path = os.path.join(current_dir, output_filename)
    
    print(f"\n💾 Сохранение результата в {output_filename}...")
    save_to_docx(merged_data, output_path)
    
    print(f"✅ Готово! Всего записей: {len(merged_data)}")
    if fallback_count > 0:
        print(f"   (Для {fallback_count} веществ данные были взяты из файла СС из-за отсутствия данных в МР)")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"❌ Критическая ошибка: {e}")
        import traceback
        traceback.print_exc()
