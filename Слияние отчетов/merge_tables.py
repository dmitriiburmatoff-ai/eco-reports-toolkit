"""
merge_reports.py

Что делает:
  1. Сканирует папку — находит до 6 файлов .docx и определяет их тип
     (МР / СГ / СС) и наличие фона.
  2. Для каждого типа склеивает разорванные таблицы-фрагменты в одну.
  3. Определяет вещества с фоном:
       МР  — ячейки cells[8]/cells[10] непусты; для кодов >= 6000
             сравнивает cells[4] по точкам между файлами.
       СГ  — ячейки cells[8]/cells[10] непусты.
       СС  — объединение веществ-с-фоном из МР и СГ,
             затем сравнение по точкам.
  4. Собирает итоговый файл: для каждого вещества блок БЕЗ фона,
     затем (если есть) блок С фоном с суффиксом в заголовке.
  5. Сохраняет: output_МР.docx, output_СГ.docx, output_СС.docx.

Использование:
  Положите скрипт в папку с файлами и запустите.
"""

import os
import sys
from copy import deepcopy
from lxml import etree
import docx
from docx.oxml import OxmlElement

WNS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
TR   = f'{{{WNS}}}tr'
TBL  = f'{{{WNS}}}tbl'
PARA = f'{{{WNS}}}p'

MARKER_MR = 'Расчет рассеивания по МРР-2017'
MARKER_SG = 'Расчет средних концентраций по МРР-2017'
MARKER_SS = 'Расчет среднесуточных концентраций'

COL_NUM  = 0
COL_CONC = 4
COL_FON  = 8
COL_FON2 = 10
COL_TYPE = 12

SUFFIX_FON = ' с учетом фоновых концентраций'


# ─────────────────────── утилиты ─────────────────────────────────────────────

def is_dash(val):
    return val.strip().replace('\xa0', '').replace(' ', '') in (
        '-', '—', '–', '', '---', '--')


def parse_float(val):
    try:
        return float(val.strip().replace(',', '.').replace(' ', '').replace('\xa0', ''))
    except Exception:
        return None


def tc_texts(tc_elem):
    """
    Читает текст ячейки с учётом <w:br/> как разделителя строк.
    Возвращает строку где <w:br/> заменён на '\n'.
    """
    parts = []
    cur = []
    for elem in tc_elem.iter():
        if elem.tag == f'{{{WNS}}}t':
            cur.append(elem.text or '')
        elif elem.tag == f'{{{WNS}}}br':
            parts.append(''.join(cur))
            cur = []
    parts.append(''.join(cur))
    return '\n'.join(parts).strip()


def row_cell_texts(tr_elem):
    """
    Возвращает список уникальных текстов ячеек строки (без дублей от merge).
    Учитывает <w:br/> как перевод строки.
    """
    cells_xml = tr_elem.findall(f'.//{{{WNS}}}tc')
    texts = []
    prev = None
    for tc in cells_xml:
        t = tc_texts(tc)
        if t != prev:
            texts.append(t)
            prev = t
    return texts


def is_data_table(table):
    """True если таблица содержит блоки веществ (строку с 'Вещество:')."""
    for row in table.rows[:30]:
        if 'Вещество:' in row.cells[0].text:
            return True
    return False


def is_substance_header(tr_elem):
    texts = row_cell_texts(tr_elem)
    return bool(texts) and 'Вещество:' in texts[0]


def parse_substance_code(tr_elem):
    """Извлекает код вещества из строки-заголовка."""
    texts = row_cell_texts(tr_elem)
    # texts[0] например: 'Вещество: 0330\nСера диоксид'
    after = texts[0].split(':', 1)[1].strip()
    # after: '0330\nСера диоксид'
    code = after.split('\n', 1)[0].strip().split()[0]
    return code


# ─────────────────────── определение типа/фона ───────────────────────────────

def get_file_type(doc):
    t0 = doc.tables[0]
    for row in t0.rows:
        text = row.cells[0].text
        if MARKER_MR in text:
            return 'МР'
        if MARKER_SG in text:
            return 'СГ'
        if MARKER_SS in text:
            return 'СС'
    return None


def file_has_background(doc):
    """
    True если в таблицах данных хоть одна строка данных
    имеет непустые ячейки Фон (cells[8] или cells[10]).
    Проверяет только таблицы с веществами.
    """
    for table in doc.tables:
        if not is_data_table(table):
            continue
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells[COL_NUM] in ('', '№') or 'Вещество:' in cells[0]:
                continue
            if len(cells) <= COL_FON2:
                continue
            if not is_dash(cells[COL_FON]) or not is_dash(cells[COL_FON2]):
                return True
    return False


# ─────────────────────── склейка таблиц ──────────────────────────────────────

def merge_tables(doc):
    """
    Склеивает все таблицы данных (с веществами) в одну.
    Возвращает объединённую таблицу.
    """
    body          = doc.element.body
    body_children = list(body)

    data_elems  = []
    data_tables = []
    for child in body_children:
        if child.tag != TBL:
            continue
        for t in doc.tables:
            if t._tbl is child and is_data_table(t):
                data_elems.append(child)
                data_tables.append(t)
                break

    if not data_tables:
        return None
    if len(data_tables) == 1:
        return data_tables[0]

    base_elem = data_elems[0]
    for extra in data_tables[1:]:
        for row in extra.rows:
            base_elem.append(deepcopy(row._tr))

    # Удаляем лишние таблицы и разделяющие параграфы между ними
    to_remove = []
    found_first = False
    for child in body_children:
        if child in data_elems:
            if not found_first:
                found_first = True
                continue
            to_remove.append(child)
        elif found_first and child.tag == PARA:
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)

    for t in doc.tables:
        if t._tbl is base_elem:
            return t
    return None


# ─────────────────────── разбор блоков веществ ───────────────────────────────

def parse_blocks(table):
    """
    Возвращает:
      header_rows — строки до первого вещества
      blocks      — dict {код: [deepcopy tr, ...]}
      order       — список кодов в порядке появления
    """
    header_rows  = []
    blocks       = {}
    order        = []
    current_code = None
    current_rows = []
    data_started = False

    for row in table.rows:
        tr = row._tr
        if is_substance_header(tr):
            if current_code is not None:
                blocks[current_code] = current_rows
            elif current_rows:
                header_rows = current_rows
            current_code = parse_substance_code(tr)
            order.append(current_code)
            current_rows = [deepcopy(tr)]
            data_started = True
        else:
            if not data_started:
                current_rows.append(deepcopy(tr))
            elif current_code is not None:
                current_rows.append(deepcopy(tr))

    if current_code is not None:
        blocks[current_code] = current_rows

    return header_rows, blocks, order


# ─────────────────────── определение веществ с фоном ─────────────────────────

def get_point_concs(blocks, code):
    """dict {номер_точки: float} — пропускает прочерки."""
    result = {}
    for tr in blocks.get(code, []):
        texts = row_cell_texts(tr)
        if len(texts) <= COL_CONC:
            continue
        num  = texts[COL_NUM].strip()
        conc = texts[COL_CONC].strip()
        if not num or num == '№' or 'Вещество:' in num:
            continue
        if is_dash(conc):
            continue
        val = parse_float(conc)
        if val is not None:
            result[num] = val
    return result


def block_has_fon_cells(blocks, code):
    """True если хоть одна строка блока имеет непустые ячейки Фон."""
    for tr in blocks.get(code, []):
        texts = row_cell_texts(tr)
        if len(texts) <= COL_FON2:
            continue
        num = texts[COL_NUM].strip()
        if not num or num == '№' or 'Вещество:' in num:
            continue
        if not is_dash(texts[COL_FON]) or not is_dash(texts[COL_FON2]):
            return True
    return False


def fon_is_greater(blocks_nofon, blocks_fon, code):
    """
    True если файл fon имеет хоть одно большее значение чем nofon.
    False если меньше (значит файлы перепутаны).
    None если нет различий.
    """
    concs_nofon = get_point_concs(blocks_nofon, code)
    concs_fon   = get_point_concs(blocks_fon,   code)
    for pt, val_nofon in concs_nofon.items():
        val_fon = concs_fon.get(pt)
        if val_fon is None:
            continue
        if val_fon > val_nofon:
            return True
        if val_fon < val_nofon:
            return False
    return None


def determine_file_order(blocks_a, blocks_b, order, known_codes=None):
    """
    Определяет какой из двух наборов блоков 'без фона', какой 'с фоном'.
    Сравнивает по точкам для первого подходящего вещества.
    Возвращает (blocks_nofon, blocks_fon, swapped: bool).
    """
    candidates = known_codes if known_codes else order
    for code in order:
        if code not in candidates:
            continue
        result = fon_is_greater(blocks_a, blocks_b, code)
        if result is True:
            return blocks_a, blocks_b, False   # B с фоном — всё верно
        if result is False:
            return blocks_b, blocks_a, True    # A с фоном — меняем
    return blocks_a, blocks_b, False  # не смогли определить — оставляем как есть


def find_fon_substances(file_type, blocks_nofon, blocks_fon, order,
                        known_fon_codes=None):
    """Возвращает set кодов веществ у которых фон реально учтён."""
    fon_codes = set()

    if file_type == 'МР':
        for code in order:
            try:
                code_int = int(code)
            except ValueError:
                code_int = 0
            if code_int >= 6000:
                if fon_is_greater(blocks_nofon, blocks_fon, code) is True:
                    fon_codes.add(code)
            else:
                if block_has_fon_cells(blocks_fon, code):
                    fon_codes.add(code)

    elif file_type == 'СГ':
        for code in order:
            if block_has_fon_cells(blocks_fon, code):
                fon_codes.add(code)

    elif file_type == 'СС':
        candidates = known_fon_codes or set()
        for code in order:
            if code not in candidates:
                continue
            if fon_is_greater(blocks_nofon, blocks_fon, code) is True:
                fon_codes.add(code)

    return fon_codes


# ─────────────────────── дописать суффикс в заголовок ────────────────────────

def append_fon_suffix(tr_elem):
    """Дописывает SUFFIX_FON в строку-заголовок вещества в том же стиле."""
    para = tr_elem.find(f'.//{{{WNS}}}p')
    if para is None:
        return
    runs = para.findall(f'{{{WNS}}}r')
    rPr_copy = None
    if runs:
        last_rPr = runs[-1].find(f'{{{WNS}}}rPr')
        if last_rPr is not None:
            rPr_copy = deepcopy(last_rPr)
    new_r = OxmlElement('w:r')
    if rPr_copy is not None:
        new_r.append(rPr_copy)
    new_t = OxmlElement('w:t')
    new_t.text = SUFFIX_FON
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    para.append(new_r)


# ─────────────────────── сборка итоговой таблицы ─────────────────────────────

def build_output(base_table, header_rows,
                 blocks_nofon, blocks_fon, order, fon_codes):
    """
    Пересобирает base_table:
    заголовок + для каждого вещества блок БЕЗ фона,
    затем (если код в fon_codes) блок С фоном + суффикс в заголовке.
    """
    tbl_elem = base_table._tbl

    for tr in tbl_elem.findall(TR):
        tbl_elem.remove(tr)

    for tr in header_rows:
        tbl_elem.append(deepcopy(tr))

    for code in order:
        for tr in blocks_nofon.get(code, []):
            tbl_elem.append(deepcopy(tr))

        if code in fon_codes and code in blocks_fon:
            fon_rows = [deepcopy(tr) for tr in blocks_fon[code]]
            if fon_rows and is_substance_header(fon_rows[0]):
                append_fon_suffix(fon_rows[0])
            for tr in fon_rows:
                tbl_elem.append(tr)

    total = len(tbl_elem.findall(TR))
    print(f"    Строк в итоговой таблице: {total}")


# ─────────────────────── обработка пары файлов ───────────────────────────────

def process_pair(file_type, paths, output_path, known_fon_codes=None):
    """
    paths — список из 1 или 2 путей.
    Возвращает set кодов веществ с фоном.
    """
    path_a = paths[0]
    path_b = paths[1] if len(paths) >= 2 else None

    print(f"\n  Файл A: {os.path.basename(path_a)}")
    doc_a = docx.Document(path_a)
    doc_b = docx.Document(path_b) if path_b else None
    if path_b:
        print(f"  Файл B: {os.path.basename(path_b)}")

    # Склейка
    print(f"  Склейка таблиц...")
    table_a = merge_tables(doc_a)
    table_b = merge_tables(doc_b) if doc_b else None

    if table_a is None:
        print(f"  ❌ Таблица данных не найдена.")
        return set()

    header_a, blocks_a, order_a = parse_blocks(table_a)
    header_b, blocks_b, order_b = (parse_blocks(table_b)
                                   if table_b else ([], {}, []))

    print(f"  Веществ в A: {len(order_a)}"
          + (f", в B: {len(order_b)}" if doc_b else ""))

    if doc_b is None:
        print(f"  Второй файл не найден — только склейка.")
        build_output(table_a, header_a, blocks_a, {}, order_a, set())
        doc_a.save(output_path)
        print(f"  ✅ Сохранено: {os.path.basename(output_path)}")
        return set()

    # Определяем какой файл с фоном, какой без
    fon_a = file_has_background(doc_a)
    fon_b = file_has_background(doc_b)
    print(f"  Фон в A: {'да' if fon_a else 'нет'}, "
          f"в B: {'да' if fon_b else 'нет'}")

    if fon_a and not fon_b:
        # A с фоном, B без
        blocks_nofon, header_nofon, order = blocks_b, header_b, order_b
        blocks_fon = blocks_a
        doc_out = doc_b; table_out = table_b
        print(f"  Порядок: A=с фоном, B=без фона")
    elif fon_b and not fon_a:
        # B с фоном, A без
        blocks_nofon, header_nofon, order = blocks_a, header_a, order_a
        blocks_fon = blocks_b
        doc_out = doc_a; table_out = table_a
        print(f"  Порядок: A=без фона, B=с фоном")
    else:
        # Оба без явного маркера (СС или МР без ячеек фона) — определяем по точкам
        blocks_nofon, blocks_fon, swapped = determine_file_order(
            blocks_a, blocks_b, order_a, known_fon_codes)
        if swapped:
            header_nofon, order = header_b, order_b
            doc_out = doc_b; table_out = table_b
            print(f"  Порядок (по точкам): A=с фоном, B=без фона")
        else:
            header_nofon, order = header_a, order_a
            doc_out = doc_a; table_out = table_a
            print(f"  Порядок (по точкам): A=без фона, B=с фоном")

    # Вещества с фоном
    fon_codes = find_fon_substances(
        file_type, blocks_nofon, blocks_fon, order, known_fon_codes)
    print(f"  Веществ с фоном ({len(fon_codes)}): {sorted(fon_codes)}")

    # Сборка и сохранение
    build_output(table_out, header_nofon,
                 blocks_nofon, blocks_fon, order, fon_codes)
    doc_out.save(output_path)
    print(f"  ✅ Сохранено: {os.path.basename(output_path)}")
    return fon_codes


# ─────────────────────── сканирование папки ──────────────────────────────────

def scan_folder(folder):
    found = {'МР': [], 'СГ': [], 'СС': []}
    for fname in sorted(os.listdir(folder)):
        if not fname.endswith('.docx'):
            continue
        if fname.startswith('~') or fname.startswith('output_'):
            continue
        fpath = os.path.join(folder, fname)
        try:
            d     = docx.Document(fpath)
            ftype = get_file_type(d)
            if ftype:
                found[ftype].append(fpath)
                print(f"  {fname} → {ftype}")
            else:
                print(f"  {fname} → тип не определён, пропущен")
        except Exception as e:
            print(f"  {fname} → ошибка: {e}")
    return found


# ─────────────────────── main ────────────────────────────────────────────────

def main():
    folder = os.path.dirname(os.path.abspath(__file__))

    print("=" * 55)
    print("Сканирование папки...")
    print("=" * 55)
    found = scan_folder(folder)

    for ftype, paths in found.items():
        n = len(paths)
        if n == 0:
            print(f"⚠ Файлы типа {ftype} не найдены.")
        elif n == 1:
            print(f"{ftype}: 1 файл (без пары — только склейка).")
        elif n == 2:
            print(f"{ftype}: 2 файла.")
        else:
            print(f"⚠ {ftype}: {n} файлов, используем первые два.")

    fon_codes_mr = set()
    if found['МР']:
        print(f"\n{'='*55}\nОбработка МР\n{'='*55}")
        fon_codes_mr = process_pair(
            'МР', found['МР'],
            os.path.join(folder, 'output_МР.docx'))

    fon_codes_sg = set()
    if found['СГ']:
        print(f"\n{'='*55}\nОбработка СГ\n{'='*55}")
        fon_codes_sg = process_pair(
            'СГ', found['СГ'],
            os.path.join(folder, 'output_СГ.docx'))

    if found['СС']:
        known_fon = fon_codes_mr | fon_codes_sg
        print(f"\n{'='*55}\nОбработка СС\n{'='*55}")
        print(f"  Вещества с фоном из МР+СГ: {sorted(known_fon)}")
        process_pair(
            'СС', found['СС'],
            os.path.join(folder, 'output_СС.docx'),
            known_fon_codes=known_fon)

    print(f"\n{'='*55}")
    print("✅ Готово.")


if __name__ == '__main__':
    main()
