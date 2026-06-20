import os
import json
import docx
from pathlib import Path
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import logging

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("process_tables.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ----------------- Загрузка справочника ПДКсс -----------------
def load_pdk_ss_dict(json_path):
    """
    Загружает output.json и извлекает словарь {код: ПДКсс (мг/м³)}.
    Если ПДКсс отсутствует или равна нулю, ключ не добавляется (или можно добавить с None).
    """
    logger.info(f"Загрузка нормативов ПДКсс из: {json_path}")
    if not json_path.exists():
        logger.error("Файл output.json не найден! Фильтрация для СГ отчётов будет отключена.")
        return {}
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    pdk_ss_dict = {}
    for item in data:
        code = item['code']
        pdk_ss_raw = item.get('pdk_ss')
        if pdk_ss_raw is not None:
            try:
                val = float(str(pdk_ss_raw).replace(',', '.'))
                # ПДКсс = 0 тоже считаем отсутствующим (не добавляем)
                if val != 0:
                    pdk_ss_dict[code] = val
            except (ValueError, AttributeError):
                logger.warning(f"Невозможно преобразовать ПДКсс для кода {code}: {pdk_ss_raw}")
    logger.info(f"Загружено {len(pdk_ss_dict)} нормативов ПДКсс (ненулевых)")
    return pdk_ss_dict

# ----------------- Определение типа отчёта -----------------
def get_report_type(doc):
    """
    Возвращает тип отчёта: 'СС', 'СГ' или 'МР'.
    Проверяет наличие ключевых фраз в тексте документа.
    """
    full_text = []
    # Собираем текст из параграфов
    for para in doc.paragraphs:
        full_text.append(para.text)
    # Также собираем текст из таблиц (на всякий случай)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    text = ' '.join(full_text)

    if 'Расчет среднесуточных концентраций' in text:
        return 'СС'
    elif 'Расчет средних концентраций по МРР-2017' in text:
        return 'СГ'
    else:
        return 'МР'

# ----------------- Парсинг документа -----------------
def extract_data_from_word(file_path, pdk_ss_dict, report_type):
    """
    Извлекает данные по веществам и точкам.
    Для отчёта СГ фильтрует вещества по наличию ненулевой ПДКсс в словаре.
    Возвращает список словарей: [{'code': код, 'name': название, 'data': {point: value}}, ...]
    """
    logger.info(f"Парсинг файла: {file_path.name}, тип отчёта: {report_type}")
    try:
        doc = docx.Document(file_path)
        results = []
        current_block = None
        processed_blocks = 0
        skipped_no_pdk = 0

        # Перебираем все таблицы и ячейки
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if 'Вещество:' in cell_text:
                        # Сохраняем предыдущий блок
                        if current_block and current_block['data']:
                            sorted_data = dict(sorted(current_block['data'].items(),
                                                      key=lambda x: int(x[0][2:])))
                            results.append({
                                'code': current_block['code'],
                                'name': current_block['name'],
                                **sorted_data
                            })
                            processed_blocks += 1

                        # Разбор заголовка вещества
                        # Формат: "Вещество: 0301 Азота диоксид ..."
                        try:
                            parts = cell_text.split(':', 1)[1].strip().split()
                        except IndexError:
                            logger.warning(f"Не удалось разделить строку: {cell_text}")
                            continue
                        if len(parts) < 2:
                            logger.warning(f"Недостаточно частей: {cell_text}")
                            current_block = None
                            continue
                        code = parts[0]
                        name = ' '.join(parts[1:])

                        # Для отчёта СГ – фильтрация по наличию ПДКсс
                        if report_type == 'СГ':
                            if code not in pdk_ss_dict:
                                logger.debug(f"Вещество {code} ({name}) пропущено (нет ПДКсс или =0)")
                                current_block = None
                                skipped_no_pdk += 1
                                continue
                            else:
                                logger.debug(f"Вещество {code} будет обработано (ПДКсс={pdk_ss_dict[code]})")

                        current_block = {
                            'code': code,
                            'name': name,
                            'data': {}
                        }
                        continue

                    if current_block:
                        # Ищем строки таблицы с номерами точек и долями ПДК
                        try:
                            cells = row.cells
                            # Проверяем наличие минимум 6 колонок (как в исходной структуре)
                            if len(cells) >= 6:
                                point_str = cells[0].text.strip()
                                try:
                                    point_num = int(point_str)
                                except ValueError:
                                    continue

                                # Доля ПДК – 5-я колонка (индекс 4)
                                dolya_str = cells[4].text.strip()
                                if dolya_str == '':
                                    dolya_str = '-'
                                current_block['data'][f'РТ{point_num}'] = dolya_str
                                logger.debug(f"  РТ{point_num}: {dolya_str}")
                        except Exception as e:
                            logger.error(f"Ошибка при обработке строки: {e}")

        # Сохранение последнего блока
        if current_block and current_block['data']:
            sorted_data = dict(sorted(current_block['data'].items(),
                                      key=lambda x: int(x[0][2:])))
            results.append({
                'code': current_block['code'],
                'name': current_block['name'],
                **sorted_data
            })
            processed_blocks += 1

        logger.info(f"Обработано веществ: {processed_blocks}, пропущено (нет ПДКсс для СГ): {skipped_no_pdk}")
        return results

    except Exception as e:
        logger.error(f"Ошибка при обработке файла {file_path.name}: {e}")
        return []

# ----------------- Создание Word-таблицы -----------------
def create_word_table(data_blocks, output_path):
    """
    Создаёт Word-документ с таблицей по извлечённым данным.
    """
    if not data_blocks:
        logger.warning("Нет данных для создания таблицы")
        return

    # Собираем все уникальные точки
    all_points = set()
    for block in data_blocks:
        for key in block.keys():
            if key.startswith('РТ'):
                all_points.add(key)
    sorted_points = sorted(all_points, key=lambda x: int(x[2:]))
    headers = ['Код', 'Название'] + sorted_points

    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Заполнение данными
    for block in data_blocks:
        row_cells = table.add_row().cells
        row_cells[0].text = block['code']
        row_cells[1].text = block['name']
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for j, point in enumerate(sorted_points, start=2):
            value = block.get(point, '')
            row_cells[j].text = str(value)
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.save(output_path)
    logger.info(f"Сохранён файл: {output_path}")

# ----------------- Основная обработка -----------------
def process_folder():
    current_dir = Path(__file__).parent
    logger.info("Начало обработки папки")

    # Загружаем справочник ПДКсс (для СГ отчётов)
    json_path = current_dir / 'output.json'
    pdk_dict = load_pdk_ss_dict(json_path)

    # Обрабатываем все DOCX файлы, начинающиеся с "output_"
    for docx_file in current_dir.glob('output_*.docx'):
        logger.info(f"--- Обработка файла: {docx_file.name} ---")
        try:
            doc = docx.Document(docx_file)
            report_type = get_report_type(doc)
            logger.info(f"Определён тип отчёта: {report_type}")
            data = extract_data_from_word(docx_file, pdk_dict, report_type)
            if data:
                out_name = docx_file.stem + '_таблица.docx'
                out_path = current_dir / out_name
                create_word_table(data, out_path)
                logger.info(f"Данные сохранены в {out_name}, веществ: {len(data)}")
            else:
                logger.warning(f"Пустые данные для файла {docx_file.name}")
        except Exception as e:
            logger.exception(f"Ошибка при обработке {docx_file.name}: {e}")

    logger.info("Обработка всех файлов завершена")

if __name__ == "__main__":
    try:
        process_folder()
    except Exception as e:
        logger.exception(f"Произошла ошибка: {e}")
