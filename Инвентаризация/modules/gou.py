# modules/gou.py
import os
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.row_generator_gou import generate_gou_rows

def set_cell_vertical_alignment(cell, align):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)
    tcPr.append(valign)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcPr.append(border)

def copy_table_preserving_all(src_table, dst_doc):
    tbl_xml = src_table._element
    new_tbl_xml = copy.deepcopy(tbl_xml)
    dst_doc._element.body.append(new_tbl_xml)
    return dst_doc.tables[-1]

def export(data, output_dir, base_name):
    rows_data = generate_gou_rows(data)
    if not rows_data:
        return

    template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'template_gou.docx')
    if not os.path.exists(template_path):
        print(f"Шаблон не найден: {template_path}")
        return

    src_doc = Document(template_path)
    src_table = src_doc.tables[0]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Результаты обследования установок очистки газа и условий их эксплуатации")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = "Times New Roman"

    table = copy_table_preserving_all(src_table, doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(8)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = style

    # Группировка по площадке и цеху
    current_plo = None
    current_cech = None
    for r in rows_data:
        plo_key = (r['plo_code'], r['plo_name'])
        cech_key = (r['cech_code'], r['cech_name'])
        if plo_key != current_plo:
            # Строка-разделитель площадки
            new_row = table.add_row()
            merged_cell = new_row.cells[0]
            for i in range(1, 11):
                merged_cell.merge(new_row.cells[i])
            merged_cell.text = f"Площадка: {r['plo_code']} {r['plo_name']}"
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(merged_cell, 'center')
            merged_cell.paragraphs[0].runs[0].bold = True
            merged_cell.paragraphs[0].style = style
            set_cell_borders(merged_cell)
            current_plo = plo_key
            current_cech = None  # сбросим, чтобы цех вывелся заново
        if cech_key != current_cech:
            # Если цех сменился, то для первой строки этого цеха нужно вывести цех,
            # а для последующих – пусто. Но в текущем генераторе cech_code и cech_name заполнены всегда.
            # Мы будем выводить их только если цех сменился, иначе пусто.
            show_cech = True
            current_cech = cech_key
        else:
            show_cech = False

        new_row = table.add_row()
        # Столбцы:
        # 1 - № цеха, 2 - Наименование цеха, 3 - № участка, 4 - источник, 5 - ГОУ, 6 - ИЗАВ,
        # 7 - эффективность проект, 8 - эффективность факт, 9 - ЗВ, 10 - коэф норм, 11 - коэф факт
        col_values = [
            r['cech_code'] if show_cech else '',
            r['cech_name'] if show_cech else '',
            r['uch_code'],
            r['source_name'],
            r['gas_name'],
            r['izav'],
            r['eff_pr'],
            r['eff_fact'],
            r['subst_name'],
            r['coef_norm'],
            r['coef_fact']
        ]
        for i, val in enumerate(col_values):
            cell = new_row.cells[i]
            cell.text = str(val) if val != '' else ''
            set_cell_vertical_alignment(cell, 'center')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].style = style
            set_cell_borders(cell)

    out_path = os.path.join(output_dir, "4.3. ГОУ.docx")
    doc.save(out_path)
    print(f"Word (ГОУ) сохранён: {out_path}")