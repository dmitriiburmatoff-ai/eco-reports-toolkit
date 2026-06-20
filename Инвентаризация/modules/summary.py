# modules/summary.py
import os
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.row_generator_summary import generate_summary_rows, format_number

def set_cell_vertical_alignment(cell, align):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)
    tcPr.append(valign)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcPr.append(border)

def copy_table_preserving_all(src_table, dst_doc):
    tbl_xml = src_table._element
    new_tbl_xml = copy.deepcopy(tbl_xml)
    dst_doc._element.body.append(new_tbl_xml)
    return dst_doc.tables[-1]

def export(data, output_dir, base_name):
    rows_data, subst_info = generate_summary_rows(data)
    if not rows_data:
        return

    template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'template_summary.docx')
    if not os.path.exists(template_path):
        print(f"Шаблон не найден: {template_path}")
        return

    src_doc = Document(template_path)
    src_table = src_doc.tables[0]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Суммарные выбросы ЗВ в атмосферный воздух, их очистка и утилизация\n(в целом по объекту ОНВ), т/год")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = "Times New Roman"

    table = copy_table_preserving_all(src_table, doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(8)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = style

    # Строка "По объекту ОНВ в целом"
    overall_row = table.add_row()
    merged_cell = overall_row.cells[0]
    for i in range(1, 10):
        merged_cell.merge(overall_row.cells[i])
    merged_cell.text = "По объекту ОНВ в целом"
    set_cell_vertical_alignment(merged_cell, 'center')
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.paragraphs[0].runs[0].bold = True
    merged_cell.paragraphs[0].style = style
    set_cell_borders(merged_cell)

    # Добавляем строки данных
    for item in rows_data:
        if item['type'] == 'header':
            new_row = table.add_row()
            merged_cell = new_row.cells[0]
            for i in range(1, 10):
                merged_cell.merge(new_row.cells[i])
            merged_cell.text = item['title']
            set_cell_vertical_alignment(merged_cell, 'center')
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            merged_cell.paragraphs[0].runs[0].bold = True
            merged_cell.paragraphs[0].style = style
            set_cell_borders(merged_cell)
        else:
            new_row = table.add_row()
            cells = new_row.cells
            cells[0].text = str(item['code'])
            cells[1].text = item['name']
            cells[2].text = format_number(item['col3'])
            cells[3].text = format_number(item['col4'])
            cells[4].text = format_number(item['col5'])
            cells[5].text = format_number(item['col6'])
            cells[6].text = format_number(item['col7'])
            cells[7].text = format_number(item['col8'])
            cells[8].text = format_number(item['col9'])
            cells[9].text = format_number(item['col10'])
            for i, cell in enumerate(cells):
                set_cell_vertical_alignment(cell, 'center')
                if i == 1:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].style = style
                set_cell_borders(cell)

    # --- Итоговые строки ---
    total_row = [0.0] * 8
    solid_row = [0.0] * 8
    liquid_row = [0.0] * 8

    for item in rows_data:
        if item['type'] == 'data':
            code = item['code']
            ag = subst_info.get(code, 1)
            vals = [item['col3'], item['col4'], item['col5'], item['col6'], item['col7'], item['col8'], item['col9'], item['col10']]
            for i, v in enumerate(vals):
                total_row[i] += v
                if ag == 1:
                    solid_row[i] += v
                else:
                    liquid_row[i] += v

    def add_summary_row(title, values):
        new_row = table.add_row()
        cells = new_row.cells
        cells[0].merge(cells[1])
        cells[0].text = title
        set_cell_vertical_alignment(cells[0], 'center')
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].style = style
        set_cell_borders(cells[0])
        for i, val in enumerate(values):
            cell = cells[i+2]
            cell.text = format_number(val)
            set_cell_vertical_alignment(cell, 'center')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].style = style
            set_cell_borders(cell)
        for cell in cells:
            set_cell_borders(cell)

    add_summary_row("Всего:", total_row)
    add_summary_row("в т. ч. твердых:", solid_row)
    add_summary_row("в т. ч. жидких и газообразных:", liquid_row)

    out_path = os.path.join(output_dir, "4.4. Суммарные.docx")
    doc.save(out_path)
    print(f"Word (Суммарные) сохранён: {out_path}")