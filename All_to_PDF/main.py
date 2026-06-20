#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
pdf_assembler.py  v2.0
══════════════════════
Сборка единого PDF-документа из разрозненных файлов проекта.

Запуск:
    python pdf_assembler.py                       # ./project_source → ./project_source/output
    python pdf_assembler.py ./src                 # ./src → ./src/output
    python pdf_assembler.py ./src ./out           # явные пути
"""

import io
import math
import os
import re
import shutil
import subprocess
import sys
import logging
import tempfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ─── Сторонние библиотеки ─────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF >= 1.23
except ImportError:
    print("ОШИБКА: pip install pymupdf")
    sys.exit(1)

try:
    from reportlab.lib.pagesizes import A4 as RL_A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.platypus import Paragraph
except ImportError:
    print("ОШИБКА: pip install reportlab")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("ОШИБКА: pip install python-docx")
    sys.exit(1)


# ═════════════════════════════════════════════════════════════════════════════
#  КОНСТАНТЫ
# ═════════════════════════════════════════════════════════════════════════════

A4_W: float = 595.276      # pt
A4_H: float = 841.890      # pt
FOOTER_H: float = 34    # 1.5 cm в pt (28.3465 pt/cm × 1.5)
PAGE_MARGIN: float = 34 # 1.5 cm отступ на титульных листах (те же 1.5 cm)
SIZE_TOL: float = 2.0      # допуск сравнения размеров, pt

# Шрифты
RL_FONT_NAME  = "TimesNewRoman"  # имя после регистрации TTF в reportlab
FITZ_FONT     = "Times-Roman"    # встроенный PDF-шрифт (fallback)

log: logging.Logger = logging.getLogger("pdf_assembler")


# ═════════════════════════════════════════════════════════════════════════════
#  УТИЛИТЫ: ЛОГИРОВАНИЕ
# ═════════════════════════════════════════════════════════════════════════════

def setup_logging(output_dir: Path) -> None:
    """Настраивает вывод в консоль (INFO) и в файл (DEBUG)."""
    log.setLevel(logging.DEBUG)
    fmt = logging.Formatter(
        "%(asctime)s  [%(levelname)-8s]  %(message)s",
        datefmt="%H:%M:%S",
    )
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    ch.setFormatter(fmt)
    log.addHandler(ch)

    log_file = output_dir / f"assembly_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    log.addHandler(fh)
    log.info("Лог: %s", log_file)


# ═════════════════════════════════════════════════════════════════════════════
#  УТИЛИТЫ: ШРИФТЫ И LIBREOFFICE
# ═════════════════════════════════════════════════════════════════════════════

def find_tnr_ttf() -> Optional[str]:
    """
    Ищет файл Times New Roman Regular (.ttf) на системе.
    Windows / Linux / macOS — проверяет стандартные пути, затем рекурсивный поиск.
    """
    candidates = [
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\Times New Roman.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
        "/usr/share/fonts/liberation/LiberationSerif-Regular.ttf",
        "/Library/Fonts/Times New Roman.ttf",
        "/System/Library/Fonts/Times.ttf",
    ]
    for p in candidates:
        if Path(p).is_file():
            log.debug("Шрифт TNR: %s", p)
            return p

    for root in [Path(r"C:\Windows\Fonts"), Path("/usr/share/fonts"),
                 Path("/Library/Fonts"), Path(os.path.expanduser("~/Library/Fonts"))]:
        if root.exists():
            for f in root.rglob("times*.ttf"):
                log.debug("Шрифт TNR (поиск): %s", f)
                return str(f)

    log.warning("Times New Roman не найден → встроенный PDF-шрифт (кириллица может отсутствовать)")
    return None


def register_rl_font(ttf_path: Optional[str]) -> str:
    """Регистрирует TTF в reportlab, возвращает имя шрифта для Canvas."""
    if ttf_path:
        try:
            pdfmetrics.registerFont(TTFont(RL_FONT_NAME, ttf_path))
            log.debug("reportlab: шрифт '%s' зарегистрирован", RL_FONT_NAME)
            return RL_FONT_NAME
        except Exception as exc:
            log.warning("Не удалось зарегистрировать TTF: %s", exc)
    return "Times-Roman"


def find_libreoffice() -> Optional[str]:
    """Ищет исполняемый файл LibreOffice (soffice) на системе."""
    candidates = [
        "soffice", "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice", "/usr/local/bin/soffice",
    ]
    for cmd in candidates:
        try:
            r = subprocess.run([cmd, "--version"], capture_output=True, timeout=5)
            if r.returncode == 0:
                log.debug("LibreOffice: %s", cmd)
                return cmd
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


# ═════════════════════════════════════════════════════════════════════════════
#  УТИЛИТЫ: PDF И ФАЙЛЫ
# ═════════════════════════════════════════════════════════════════════════════

def page_count(pdf_path: str) -> int:
    """Количество страниц в PDF."""
    try:
        d = fitz.open(pdf_path)
        n = len(d)
        d.close()
        return n
    except Exception:
        return 0


def merge_pdfs(pdf_paths: List[str], output: str) -> bool:
    """Объединяет PDF-файлы в один в заданном порядке."""
    try:
        dst = fitz.open()
        for p in pdf_paths:
            if not Path(p).exists():
                log.warning("  merge: не найден → %s", p)
                continue
            src = fitz.open(p)
            dst.insert_pdf(src)
            src.close()
        Path(output).parent.mkdir(parents=True, exist_ok=True)
        dst.save(output, garbage=4, deflate=True)
        dst.close()
        return True
    except Exception as exc:
        log.error("Ошибка слияния PDF: %s", exc)
        return False


def make_blank_pdf() -> bytes:
    """Создаёт PDF с одной пустой страницей A4."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=RL_A4)
    c.showPage()
    c.save()
    return buf.getvalue()


def sort_key_numeric(name: str) -> Tuple[int, int]:
    """
    Ключ числовой сортировки имён файлов.
    '01_...' → (1,0) | '02.1_...' → (2,1) | '02.2_...' → (2,2)
    """
    m = re.match(r"^(\d+)(?:\.(\d+))?", name)
    return (int(m.group(1)), int(m.group(2) or 0)) if m else (999, 0)


# ═════════════════════════════════════════════════════════════════════════════
#  ИСПРАВЛЕНИЕ 1: Авто-определение главного документа
# ═════════════════════════════════════════════════════════════════════════════

def find_main_document(source_dir: Path) -> Optional[Path]:
    """
    Находит основной документ Word в корне папки источника.

    Алгоритм:
        1. Ищет все .docx в корне (не рекурсивно).
        2. Если один — возвращает его.
        3. Если несколько — возвращает тот, что содержит таблицу приложений
           (3+ столбца, «Приложение» в первом столбце).
        4. Если не нашёл по критерию — первый по имени.
    """
    docx_files = sorted(
        [f for f in source_dir.iterdir() if f.suffix.lower() == ".docx" and f.is_file()]
    )

    if not docx_files:
        log.error("Нет .docx файлов в корне: %s", source_dir)
        return None

    if len(docx_files) == 1:
        log.info("  Основной документ: %s", docx_files[0].name)
        return docx_files[0]

    # Ищем файл с таблицей приложений
    for f in docx_files:
        try:
            doc = Document(str(f))
            for table in doc.tables:
                if len(table.columns) >= 3:
                    if any("приложение" in r.cells[0].text.lower()
                           for r in table.rows):
                        log.info(
                            "  Основной документ (по таблице): %s "
                            "(из %d .docx в папке)",
                            f.name, len(docx_files),
                        )
                        return f
        except Exception:
            continue

    log.warning(
        "  Таблица приложений не найдена ни в одном .docx. "
        "Используется первый файл: %s",
        docx_files[0].name,
    )
    return docx_files[0]


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 1: parse_main_document
# ═════════════════════════════════════════════════════════════════════════════

def parse_main_document(docx_path: str) -> Tuple[List[Dict], List[Dict]]:
    """
    Разбирает основной документ Word.

    Возвращает:
        app_list — [{"number": str, "name": str, "row_idx": int}]
        headings — [{"level": int, "title": str}]
    """
    log.info("━" * 56)
    log.info("ЭТАП 1/9  Разбор: %s", Path(docx_path).name)

    doc = Document(docx_path)

    # Заголовки H1–H3 для закладок
    headings: List[Dict] = []
    for para in doc.paragraphs:
        style = para.style.name.lower()
        for lvl in (1, 2, 3):
            if f"heading {lvl}" in style and para.text.strip():
                headings.append({"level": lvl, "title": para.text.strip()})
                break

    log.info("  Заголовков H1–H3: %d", len(headings))
    for h in headings:
        log.debug("  H%d: %s", h["level"], h["title"][:60])

    # Таблица приложений
    app_list: List[Dict] = []
    for table in doc.tables:
        if len(table.columns) < 3:
            continue
        if not any("приложение" in r.cells[0].text.lower() for r in table.rows):
            continue

        for row_idx, row in enumerate(table.rows):
            col0 = row.cells[0].text.strip()
            col1 = row.cells[1].text.strip()
            if "приложение" in col0.lower() and col1:
                app_list.append({
                    "number":  col0,
                    "name":    col1,
                    "row_idx": row_idx,
                })
                log.debug("  [%d] %s → %s", row_idx, col0, col1)
        break

    if not app_list:
        log.warning("  Таблица приложений не найдена или пуста")
    else:
        log.info("  Приложений в TOC: %d", len(app_list))

    return app_list, headings


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 2: convert_to_pdf
# ═════════════════════════════════════════════════════════════════════════════

def convert_to_pdf(src: str, dst: str, lo_cmd: str) -> bool:
    """
    Конвертирует .docx / .xlsx → PDF через LibreOffice headless.
    .pdf копируется без изменений.
    """
    ext = Path(src).suffix.lower()
    log.debug("  convert: %s", Path(src).name)

    if ext == ".pdf":
        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        return True

    supported = {".docx", ".doc", ".odt", ".xlsx", ".xls", ".ods"}
    if ext not in supported:
        log.error("  Неподдерживаемый формат: %s", ext)
        return False

    with tempfile.TemporaryDirectory() as td:
        try:
            r = subprocess.run(
                [lo_cmd, "--headless", "--convert-to", "pdf",
                 "--outdir", td, str(Path(src).resolve())],
                capture_output=True, text=True, timeout=180,
            )
            if r.returncode != 0:
                log.error("  LO ошибка (код %d): %s",
                          r.returncode, r.stderr[:300])
                return False

            expected = Path(td) / (Path(src).stem + ".pdf")
            if not expected.exists():
                found = list(Path(td).glob("*.pdf"))
                if not found:
                    log.error("  LO не создал PDF: %s", Path(src).name)
                    return False
                expected = found[0]

            Path(dst).parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(expected, dst)
            return True

        except subprocess.TimeoutExpired:
            log.error("  Таймаут (180 с): %s", Path(src).name)
            return False
        except Exception as exc:
            log.error("  Ошибка конвертации %s: %s", Path(src).name, exc)
            return False


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 3: create_title_page  (ИСПРАВЛЕНИЕ 4 — отступы + перенос строк)
# ═════════════════════════════════════════════════════════════════════════════

def create_title_page(
    app_number: str,
    app_name: str,
    output: str,
    rl_font: str,
) -> bool:
    """
    Создаёт PDF-страницу титульного листа приложения.

    Формат:
        - Times New Roman (или fallback), 16 pt
        - Отступы 1.5 cm с каждой стороны
        - Абзац с переносом строк в пределах рабочей ширины
        - Вертикальное и горизонтальное центрирование
        - Строка 1: номер приложения
        - Строка 2: название приложения
    """
    log.debug("  Титул: %s", app_number)
    try:
        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=RL_A4)
        page_w, page_h = RL_A4

        margin   = PAGE_MARGIN            # 1.5 cm
        avail_w  = page_w - 2 * margin   # рабочая ширина
        gap      = 14                     # pt между блоками

        # Стиль: Times New Roman, 16 pt, центрирование, межстрочный интервал 1.4
        style = ParagraphStyle(
            name="AppTitle",
            fontName=rl_font,
            fontSize=16,
            leading=22,       # межстрочный интервал
            alignment=1,      # TA_CENTER
            wordWrap="CJK",   # надёжный перенос для кирилицы
        )

        # Экранируем HTML-спецсимволы для Paragraph
        def esc(s: str) -> str:
            return (s.replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;")
                     .replace("\n", "<br/>"))

        p_num  = Paragraph(esc(app_number), style)
        p_name = Paragraph(esc(app_name),   style)

        # Измеряем итоговую высоту после переноса строк
        _, h_num  = p_num.wrap(avail_w, page_h)
        _, h_name = p_name.wrap(avail_w, page_h)

        total_h = h_num + gap + h_name

        # Вертикальный центр страницы
        # ReportLab: Y идёт снизу вверх, drawOn(canvas, x, y_bottom)
        y_bottom_name = page_h / 2 - total_h / 2
        y_bottom_num  = y_bottom_name + h_name + gap

        p_num.drawOn(c,  margin, y_bottom_num)
        p_name.drawOn(c, margin, y_bottom_name)

        c.showPage()
        c.save()

        Path(output).parent.mkdir(parents=True, exist_ok=True)
        Path(output).write_bytes(buf.getvalue())
        return True

    except Exception as exc:
        log.error("  Ошибка создания титула (%s): %s", app_number, exc)
        return False


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 4: scale_to_a4  (ИСПРАВЛЕНИЕ 3 — сохранение ориентации)
# ═════════════════════════════════════════════════════════════════════════════

def scale_to_a4(src: str, dst: str) -> bool:
    """
    Масштабирует страницы PDF до A4, сохраняя ориентацию каждого листа.

    Логика:
        • Определяем ориентацию исходной страницы (книжная / альбомная).
        • Целевой размер:
            - Книжная (h >= w)  → A4 книжный  (595 × 842 pt)
            - Альбомная (w > h) → A4 альбомный (842 × 595 pt)
        • Если страница уже точно соответствует целевому A4 → копируем без изменений.
        • Иначе → вписываем контент в целевой A4 с сохранением пропорций
          и центрированием (векторное качество через show_pdf_page).
    """
    log.info("━" * 56)
    log.info("ЭТАП 4  Масштабирование до A4: %s", Path(src).name)

    try:
        src_doc = fitz.open(src)
        dst_doc = fitz.open()
        total   = len(src_doc)
        scaled  = 0

        for idx in range(total):
            page = src_doc[idx]
            w, h = page.rect.width, page.rect.height

            # ── Определяем целевую ориентацию ─────────────────────────────
            if w > h:
                # Исходная альбомная → A4 альбомный
                target_w, target_h = A4_H, A4_W   # 841.89 × 595.28
            else:
                # Исходная книжная → A4 книжный
                target_w, target_h = A4_W, A4_H   # 595.28 × 841.89

            # Проверяем: уже точно A4 нужной ориентации?
            already_a4 = (
                abs(w - target_w) < SIZE_TOL and
                abs(h - target_h) < SIZE_TOL
            )

            new_p = dst_doc.new_page(width=target_w, height=target_h)

            if already_a4:
                new_p.show_pdf_page(new_p.rect, src_doc, idx)
            else:
                # Вписываем в целевой размер с сохранением пропорций
                ratio = min(target_w / w, target_h / h)
                nw, nh = w * ratio, h * ratio
                ox = (target_w - nw) / 2
                oy = (target_h - nh) / 2

                new_p.show_pdf_page(
                    fitz.Rect(ox, oy, ox + nw, oy + nh),
                    src_doc,
                    idx,
                )
                scaled += 1
                log.debug(
                    "  стр.%d: %.0f×%.0f → %.0f×%.0f  (ratio=%.3f, %s)",
                    idx + 1, w, h, target_w, target_h, ratio,
                    "альбом→A4L" if w > h else "книга→A4P",
                )

        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        dst_doc.save(dst, garbage=4, deflate=True)
        dst_doc.close()
        src_doc.close()

        if scaled:
            log.info("  Масштабировано: %d из %d стр.", scaled, total)
        else:
            log.info("  Все %d стр. уже A4", total)
        return True

    except Exception as exc:
        log.error("  Ошибка масштабирования: %s", exc)
        return False


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 5: add_page_numbers  (ИСПРАВЛЕНИЕ 5 — альбом: номер слева)
# ═════════════════════════════════════════════════════════════════════════════

def add_page_numbers(
    src: str,
    dst: str,
    tnr_ttf: Optional[str],
) -> Dict[int, int]:
    """
    Нумерует страницы PDF.

    Правила:
        • Стр. 1 (index=0) — без номера (титульный лист проекта).
        • Стр. 2+ → номер = index + 1.
        • Шрифт: Times New Roman 12 pt, колонтитул 1.5 cm.

    Книжная страница:
        Номер внизу по центру (rect: x=0..w, y=h-FOOTER_H..h).

    Альбомная страница (w/h > 1.15):
        Номер у ЛЕВОГО края, по центру высоты, текст повёрнут на 90° CCW.
        При физическом повороте листа CCW для чтения → номер оказывается
        внизу по центру читаемой страницы.

        Позиционирование (координаты fitz, Y↓):
            x = FOOTER_H / 2               ← центр полосы у левого края
            y = h/2 + text_half_width       ← якорь ниже центра
        rotate=90 (CCW) в fitz: текст идёт вверх от якоря.
        Итог: текст центрирован по вертикали на высоте h/2. ✓

    Возвращает:
        {page_index_0based: displayed_number} — пустой dict при ошибке.
    """
    log.info("━" * 56)
    log.info("ЭТАП 5  Нумерация: %s", Path(src).name)

    # Параметры шрифта для PyMuPDF
    fkw: Dict = {"fontsize": 12, "color": (0, 0, 0)}
    if tnr_ttf and Path(tnr_ttf).exists():
        fkw["fontfile"] = tnr_ttf
        fkw["fontname"] = "TNR"
    else:
        fkw["fontname"] = FITZ_FONT

    num_map: Dict[int, int] = {}

    try:
        doc = fitz.open(src)
        total = len(doc)

        for idx in range(total):
            if idx == 0:
                continue  # титульный лист проекта — без номера

            disp_num = idx + 1
            num_str  = str(disp_num)
            num_map[idx] = disp_num

            page = doc[idx]
            w, h = page.rect.width, page.rect.height
            is_landscape = (h > 0) and (w / h > 1.15)

            if is_landscape:
                # ── Альбомная: номер у ЛЕВОГО края ────────────────────────
                # fitz rotate=90 (CCW): текст идёт вверх (y уменьшается)
                # Якорная точка: x = центр полосы FOOTER_H, y = ниже центра
                # Грубая оценка ширины строки (шрифт моноширинный ~0.55 em)
                text_half_w = fkw["fontsize"] * len(num_str) * 0.30
                x_anchor    = FOOTER_H / 2
                y_anchor    = h / 2 + text_half_w

                page.insert_text(
                    fitz.Point(x_anchor, y_anchor),
                    num_str,
                    rotate=270,
                    **fkw,
                )
                log.debug(
                    "  стр.%d (альбом %.0f×%.0f) → «%s» [левый край]",
                    idx + 1, w, h, num_str,
                )
            else:
                # ── Книжная: номер внизу по центру ────────────────────────
                footer_rect = fitz.Rect(0.0, h - FOOTER_H, w, h)
                page.insert_textbox(
                    footer_rect,
                    num_str,
                    align=fitz.TEXT_ALIGN_CENTER,
                    **fkw,
                )
                log.debug("  стр.%d → «%s»", idx + 1, num_str)

        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        doc.save(dst, garbage=4, deflate=True)
        doc.close()

        log.info("  Пронумеровано %d стр. (из %d)", len(num_map), total)
        return num_map

    except Exception as exc:
        log.error("  Ошибка нумерации: %s", exc)
        return {}


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 6: update_table_of_apps
# ═════════════════════════════════════════════════════════════════════════════

def update_table_of_apps(
    src_docx: str,
    dst_docx: str,
    row_to_page: Dict[int, int],
) -> bool:
    """
    Вписывает номера страниц в третий столбец таблицы содержания приложений.

    Args:
        src_docx:    исходный .docx
        dst_docx:    путь для обновлённого .docx
        row_to_page: {row_idx_в_таблице: displayed_page_number}
    """
    log.info("━" * 56)
    log.info("ЭТАП 6  Обновление TOC в документе")

    try:
        doc = Document(src_docx)

        target_table = None
        for table in doc.tables:
            if len(table.columns) >= 3:
                if any("приложение" in r.cells[0].text.lower() for r in table.rows):
                    target_table = table
                    break

        if target_table is None:
            log.warning("  Таблица не найдена → пропуск обновления")
            return False

        updated = 0
        for row_idx, page_num in row_to_page.items():
            if row_idx >= len(target_table.rows):
                continue
            cell = target_table.rows[row_idx].cells[2]

            # Очищаем ячейку
            for para in cell.paragraphs[1:]:
                para._element.getparent().remove(para._element)
            cell.paragraphs[0].clear()

            run = cell.paragraphs[0].add_run(str(page_num))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            updated += 1
            log.debug("  row %d → стр. %d", row_idx, page_num)

        Path(dst_docx).parent.mkdir(parents=True, exist_ok=True)
        doc.save(dst_docx)
        log.info("  Обновлено строк: %d", updated)
        return True

    except Exception as exc:
        log.error("  Ошибка обновления TOC: %s", exc)
        return False


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 7: insert_blank_pages
# ═════════════════════════════════════════════════════════════════════════════

def insert_blank_pages(
    src: str,
    dst: str,
    title_pages: List[Dict],
) -> Tuple[bool, List[Dict]]:
    """
    Вставляет пустые страницы для двусторонней печати ПОСЛЕ нумерации.

    Пустые страницы:
        • Не имеют напечатанного номера.
        • Не учитываются в нумерации и в TOC.
        • Используются только для выравнивания при брошюровке.

    Правила:
        1. ПЕРЕД каждым титулом: если displayed_num нечётный → пустая страница.
        2. ПОСЛЕ каждого титула: всегда пустая страница.

    Args:
        title_pages: [{"index": int (0-based в src), "displayed_num": int,
                        "app_name": str}]

    Returns:
        (success, updated_title_pages)  — updated содержит новые физ. индексы.
    """
    log.info("━" * 56)
    log.info("ЭТАП 7  Вставка пустых страниц")

    try:
        src_doc   = fitz.open(src)
        blank_doc = fitz.open("pdf", make_blank_pdf())
        dst_doc   = fitz.open()

        title_map = {t["index"]: t for t in title_pages}
        new_titles: List[Dict] = []
        physical  = 0
        inserted  = 0

        for old_idx in range(len(src_doc)):
            if old_idx in title_map:
                tinfo = title_map[old_idx]
                disp  = tinfo["displayed_num"]

                # Правило 1: нечётный displayed_num → пустая страница перед титулом
                if disp % 2 == 1:
                    dst_doc.insert_pdf(blank_doc)
                    physical += 1
                    inserted += 1
                    log.debug(
                        "  [blank] ПЕРЕД «%s» (стр.%d нечётная)",
                        tinfo["app_name"][:40], disp,
                    )

                # Запоминаем НОВУЮ физическую позицию
                new_titles.append({
                    "index":         physical,
                    "displayed_num": disp,
                    "app_name":      tinfo["app_name"],
                })

            # Копируем оригинальную страницу
            dst_doc.insert_pdf(src_doc, from_page=old_idx, to_page=old_idx)
            physical += 1

            # Правило 2: пустая страница после каждого титула
            if old_idx in title_map:
                dst_doc.insert_pdf(blank_doc)
                physical += 1
                inserted += 1
                log.debug(
                    "  [blank] ПОСЛЕ «%s»",
                    title_map[old_idx]["app_name"][:40],
                )

        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        dst_doc.save(dst, garbage=4, deflate=True)
        dst_doc.close()
        src_doc.close()
        blank_doc.close()

        log.info(
            "  Вставлено пустых: %d | Итого стр.: %d",
            inserted, physical,
        )
        return True, new_titles

    except Exception as exc:
        log.error("  Ошибка вставки пустых страниц: %s", exc)
        return False, title_pages


# ═════════════════════════════════════════════════════════════════════════════
#  ИСПРАВЛЕНИЕ 2: Поиск заголовков в сгенерированном PDF
# ═════════════════════════════════════════════════════════════════════════════

def find_heading_pages_in_pdf(
    pdf_path: str,
    headings: List[Dict],
) -> List[Dict]:
    """
    Ищет физические страницы, на которых расположены заголовки.

    Для каждого заголовка перебирает страницы PDF и ищет полное
    совпадение текста. Первое вхождение считается позицией заголовка.

    Args:
        pdf_path: путь к PDF (после вставки пустых страниц — финальная структура)
        headings: [{"level": int, "title": str}]

    Returns:
        [{"level": int, "title": str, "page_1based": int}]
        Заголовки без найденной страницы пропускаются.
    """
    log.info("━" * 56)
    log.info("  Поиск позиций заголовков в PDF")

    result: List[Dict] = []
    try:
        doc = fitz.open(pdf_path)
        total = len(doc)

        # Кэшируем текст страниц (не держим всё в памяти сразу)
        for h in headings:
            title = h["title"].strip()
            if not title:
                continue

            found_page = None
            for idx in range(total):
                page_text = doc[idx].get_text("text")
                if title in page_text:
                    found_page = idx + 1  # 1-based
                    break

            if found_page is not None:
                result.append({
                    "level":      h["level"],
                    "title":      h["title"],
                    "page_1based": found_page,
                })
                log.debug(
                    "  H%d «%s» → стр. %d",
                    h["level"], title[:50], found_page,
                )
            else:
                log.debug("  H%d «%s» — не найден в PDF", h["level"], title[:50])

        doc.close()
        log.info("  Найдено заголовков: %d из %d", len(result), len(headings))

    except Exception as exc:
        log.error("  Ошибка поиска заголовков: %s", exc)

    return result


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 8: add_bookmarks
# ═════════════════════════════════════════════════════════════════════════════

def add_bookmarks(
    src: str,
    dst: str,
    heading_bookmarks: List[Dict],
    title_bookmarks: List[Dict],
) -> bool:
    """
    Добавляет PDF-закладки (Outlines).

    heading_bookmarks: [{"level": int, "title": str, "page_1based": int}]
    title_bookmarks:   [{"title": str, "page_1based": int}]

    Закладки из заголовков сохраняют иерархию (H1→H2→H3).
    Закладки из титульных листов — уровень 1.
    """
    log.info("━" * 56)
    log.info("ЭТАП 8  Добавление закладок")

    try:
        doc = fitz.open(src)
        toc: List[List] = []

        for h in heading_bookmarks:
            if h.get("page_1based"):
                toc.append([h["level"], h["title"], h["page_1based"]])

        for t in title_bookmarks:
            if t.get("page_1based"):
                toc.append([1, t["title"], t["page_1based"]])

        # Сортируем по странице (PyMuPDF требует монотонного порядка)
        toc.sort(key=lambda x: x[2])

        doc.set_toc(toc)
        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        doc.save(dst, garbage=4, deflate=True)
        doc.close()

        log.info("  Закладок: %d", len(toc))
        for t in toc:
            log.debug("  L%d стр.%d — %s", t[0], t[2], t[1][:60])
        return True

    except Exception as exc:
        log.error("  Ошибка закладок: %s", exc)
        return False


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 9: add_last_page
# ═════════════════════════════════════════════════════════════════════════════

def add_last_page(src: str, dst: str, rl_font: str) -> bool:
    """
    Добавляет информационную страницу в конец документа.

    Содержит: общее число страниц, листов для печати (ceil / 2), дату.
    """
    log.info("━" * 56)
    log.info("ЭТАП 9  Добавление последней страницы")

    try:
        base = fitz.open(src)
        total  = len(base)
        sheets = math.ceil(total / 2)
        base.close()

        buf = io.BytesIO()
        c   = rl_canvas.Canvas(buf, pagesize=RL_A4)
        pw, ph = RL_A4
        cx = pw / 2

        c.setFont(rl_font, 14)
        c.drawCentredString(cx, ph / 2 + 68, "Сведения о документе")

        c.setFont(rl_font, 12)
        lines = [
            f"Всего страниц в PDF:                   {total}",
            f"Листов для двусторонней печати:    {sheets}",
            f"Сформирован:  {datetime.now().strftime('%d.%m.%Y  %H:%M')}",
        ]
        y = ph / 2 + 28
        for line in lines:
            c.drawCentredString(cx, y, line)
            y -= 26

        c.showPage()
        c.save()

        main_doc  = fitz.open(src)
        extra_doc = fitz.open("pdf", buf.getvalue())
        main_doc.insert_pdf(extra_doc)

        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        main_doc.save(dst, garbage=4, deflate=True)
        main_doc.close()
        extra_doc.close()

        log.info(
            "  Последняя стр. добавлена. Итого: %d стр., %d листов",
            total + 1, math.ceil((total + 1) / 2),
        )
        return True

    except Exception as exc:
        log.error("  Ошибка последней страницы: %s", exc)
        return False


# ═════════════════════════════════════════════════════════════════════════════
#  МОДУЛЬ 10: main_pipeline
# ═════════════════════════════════════════════════════════════════════════════

def main_pipeline(source_dir: str, output_dir: str) -> bool:
    """
    Оркестратор сборки PDF.

    ╔══════════════════════════════════════════════════════╗
    ║  Pass 1 — черновик (определение нумерации страниц)  ║
    ║    convert → assemble → scale → number              ║
    ║    → записать displayed_num каждого титульного листа ║
    ╠══════════════════════════════════════════════════════╣
    ║  Pass 2 — финал (обновлённый TOC)                   ║
    ║    update TOC → re-convert main.docx                ║
    ║    → assemble → scale → number                      ║
    ║    → insert blank pages (ПОСЛЕ нумерации)           ║
    ║    → find heading pages → bookmarks → last page     ║
    ╚══════════════════════════════════════════════════════╝

    Пустые страницы не влияют на displayed_num и записи в TOC.
    """
    src_root = Path(source_dir).resolve()
    out_root = Path(output_dir).resolve()
    out_root.mkdir(parents=True, exist_ok=True)

    setup_logging(out_root)

    log.info("═" * 56)
    log.info("  PDF ASSEMBLY PIPELINE  v2.0")
    log.info("  Источник : %s", src_root)
    log.info("  Результат: %s", out_root)
    log.info("═" * 56)

    # ── Системные ресурсы ──────────────────────────────────────────────────
    lo_cmd = find_libreoffice()
    if not lo_cmd:
        log.error(
            "LibreOffice не найден!\n"
            "  Ubuntu : sudo apt-get install libreoffice\n"
            "  macOS  : brew install --cask libreoffice\n"
            "  Windows: https://www.libreoffice.org/download/download/"
        )
        return False
    log.info("LibreOffice : %s", lo_cmd)

    tnr_ttf = find_tnr_ttf()
    rl_font = register_rl_font(tnr_ttf)
    log.info("Шрифт RL    : %s", rl_font)
    log.info("Шрифт TTF   : %s", tnr_ttf or "встроенный")

    tmpdir = Path(tempfile.mkdtemp(prefix="pdf_asm_"))
    log.debug("Temp: %s", tmpdir)

    try:
        # ── Проверка структуры ─────────────────────────────────────────────
        apps_dir = src_root / "applications"
        if not apps_dir.exists():
            log.error("Папка applications/ не найдена: %s", apps_dir)
            return False

        # ИСПРАВЛЕНИЕ 1: авто-определение главного документа
        main_docx_path = find_main_document(src_root)
        if main_docx_path is None:
            return False
        log.info("Главный документ: %s", main_docx_path.name)

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 1: Разбор
        # ══════════════════════════════════════════════════════════════════
        app_list, headings = parse_main_document(str(main_docx_path))

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 2: Конвертация
        # ══════════════════════════════════════════════════════════════════
        log.info("━" * 56)
        log.info("ЭТАП 2/9  Конвертация в PDF")

        main_pdf_orig = tmpdir / "main_orig.pdf"
        if not convert_to_pdf(str(main_docx_path), str(main_pdf_orig), lo_cmd):
            log.error("Критично: не удалось конвертировать главный документ")
            return False
        log.info("  [OK] %s → PDF (%d стр.)",
                 main_docx_path.name, page_count(str(main_pdf_orig)))

        SUPPORTED = {".docx", ".doc", ".odt", ".xlsx", ".xls", ".ods", ".pdf"}
        raw_files = sorted(
            [f for f in apps_dir.iterdir()
             if f.suffix.lower() in SUPPORTED and f.is_file()],
            key=lambda f: sort_key_numeric(f.name),
        )
        log.info("  Файлов приложений: %d", len(raw_files))

        # converted_map: str(original_path) → str(pdf_path)
        converted_map: Dict[str, str] = {}
        for f in raw_files:
            out_pdf = tmpdir / f"conv_{f.stem}.pdf"
            if convert_to_pdf(str(f), str(out_pdf), lo_cmd):
                converted_map[str(f)] = str(out_pdf)
                log.info("  [OK] %s → PDF (%d стр.)",
                         f.name, page_count(str(out_pdf)))
            else:
                log.warning("  [SKIP] %s", f.name)

        # Группируем по major-номеру файла
        app_groups: Dict[int, List[str]] = defaultdict(list)
        for f in raw_files:
            major, _ = sort_key_numeric(f.name)
            if str(f) in converted_map:
                app_groups[major].append(converted_map[str(f)])

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 3: Титульные листы
        # ══════════════════════════════════════════════════════════════════
        log.info("━" * 56)
        log.info("ЭТАП 3/9  Создание титульных листов")

        title_pdfs: Dict[int, str] = {}
        for i, app in enumerate(app_list):
            major = i + 1
            out   = tmpdir / f"title_{major:03d}.pdf"
            if create_title_page(app["number"], app["name"], str(out), rl_font):
                title_pdfs[major] = str(out)
                log.info("  [OK] %s — %s", app["number"], app["name"])
            else:
                log.error("  [FAIL] %s", app["number"])

        # ══════════════════════════════════════════════════════════════════
        #  PASS 1: Черновая сборка → нумерация → фиксируем номера титулов
        # ══════════════════════════════════════════════════════════════════
        log.info("━" * 56)
        log.info("PASS 1  Черновая сборка для определения нумерации")

        ordered_v1: List[str] = [str(main_pdf_orig)]
        title_idx_v1: Dict[int, int] = {}  # major → 0-based idx
        cursor = page_count(str(main_pdf_orig))

        for i, app in enumerate(app_list):
            major = i + 1
            if major in title_pdfs:
                title_idx_v1[major] = cursor
                ordered_v1.append(title_pdfs[major])
                cursor += 1
            for pdf in app_groups.get(major, []):
                ordered_v1.append(pdf)
                cursor += page_count(pdf)

        assembled_v1 = tmpdir / "assembled_v1.pdf"
        if not merge_pdfs(ordered_v1, str(assembled_v1)):
            return False
        log.info("  Черновик: %d стр.", page_count(str(assembled_v1)))

        scaled_v1   = tmpdir / "scaled_v1.pdf"
        numbered_v1 = tmpdir / "numbered_v1.pdf"

        if not scale_to_a4(str(assembled_v1), str(scaled_v1)):
            return False
        num_map_v1 = add_page_numbers(str(scaled_v1), str(numbered_v1), tnr_ttf)
        if not num_map_v1:
            log.error("Нумерация Pass 1 не удалась")
            return False

        # Фиксируем displayed_num для каждого титула
        log.info("  Номера страниц титульных листов (Pass 1):")
        row_to_page: Dict[int, int] = {}
        title_info:  List[Dict]     = []

        for i, app in enumerate(app_list):
            major = i + 1
            if major not in title_idx_v1:
                continue
            old_idx  = title_idx_v1[major]
            disp_num = num_map_v1.get(old_idx, old_idx + 1)

            row_to_page[app["row_idx"]] = disp_num
            title_info.append({
                "major":         major,
                "row_idx":       app["row_idx"],
                "index_v1":      old_idx,
                "displayed_num": disp_num,
                "app_name":      f"{app['number']} — {app['name']}",
            })
            log.info(
                "  [%d] %s → стр. %d (физ.idx %d)",
                major, app["number"], disp_num, old_idx,
            )

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 6: Обновление TOC + перегенерация главного документа
        # ══════════════════════════════════════════════════════════════════
        updated_docx = tmpdir / "main_updated.docx"
        toc_ok = update_table_of_apps(
            str(main_docx_path), str(updated_docx), row_to_page,
        )
        src_for_v2 = str(updated_docx) if toc_ok else str(main_docx_path)

        main_pdf_v2 = tmpdir / "main_v2.pdf"
        if not convert_to_pdf(src_for_v2, str(main_pdf_v2), lo_cmd):
            log.warning("  Перегенерация не удалась → используется оригинал")
            main_pdf_v2 = main_pdf_orig

        # ══════════════════════════════════════════════════════════════════
        #  PASS 2: Финальная сборка с обновлённым TOC
        # ══════════════════════════════════════════════════════════════════
        log.info("━" * 56)
        log.info("PASS 2  Финальная сборка (обновлённый TOC)")

        ordered_v2: List[str]     = [str(main_pdf_v2)]
        title_idx_v2: Dict[int, int] = {}
        cursor = page_count(str(main_pdf_v2))

        for i, app in enumerate(app_list):
            major = i + 1
            if major in title_pdfs:
                title_idx_v2[major] = cursor
                ordered_v2.append(title_pdfs[major])
                cursor += 1
            for pdf in app_groups.get(major, []):
                ordered_v2.append(pdf)
                cursor += page_count(pdf)

        assembled_v2 = tmpdir / "assembled_v2.pdf"
        if not merge_pdfs(ordered_v2, str(assembled_v2)):
            return False
        log.info("  Финальный черновик: %d стр.", page_count(str(assembled_v2)))

        scaled_v2   = tmpdir / "scaled_v2.pdf"
        numbered_v2 = tmpdir / "numbered_v2.pdf"

        if not scale_to_a4(str(assembled_v2), str(scaled_v2)):
            return False
        num_map_v2 = add_page_numbers(str(scaled_v2), str(numbered_v2), tnr_ttf)
        if not num_map_v2:
            log.error("Нумерация Pass 2 не удалась")
            return False

        # Список титулов для вставки пустых страниц (позиции из Pass 2)
        titles_for_blanks: List[Dict] = []
        for ti in title_info:
            major  = ti["major"]
            idx_v2 = title_idx_v2.get(major)
            if idx_v2 is None:
                continue
            disp = num_map_v2.get(idx_v2, ti["displayed_num"])
            if disp != ti["displayed_num"]:
                log.debug(
                    "  [%d] displayed уточнён Pass1→Pass2: %d→%d",
                    major, ti["displayed_num"], disp,
                )
            titles_for_blanks.append({
                "index":         idx_v2,
                "displayed_num": disp,
                "app_name":      ti["app_name"],
            })

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 7: Вставка пустых страниц (ПОСЛЕ нумерации)
        # ══════════════════════════════════════════════════════════════════
        with_blanks = tmpdir / "with_blanks.pdf"
        ok, updated_titles = insert_blank_pages(
            str(numbered_v2), str(with_blanks), titles_for_blanks,
        )
        if not ok:
            log.warning("  Пустые страницы не вставлены → продолжаем без них")
            with_blanks    = numbered_v2
            updated_titles = titles_for_blanks

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 8: Закладки
        #  ИСПРАВЛЕНИЕ 2: ищем реальные страницы заголовков в финальном PDF
        # ══════════════════════════════════════════════════════════════════
        heading_bm = find_heading_pages_in_pdf(str(with_blanks), headings)

        title_bm = [
            {
                "title":       t["app_name"],
                "page_1based": t["index"] + 1,   # 0-based → 1-based
            }
            for t in updated_titles
        ]

        with_bm = tmpdir / "with_bookmarks.pdf"
        if not add_bookmarks(str(with_blanks), str(with_bm), heading_bm, title_bm):
            log.warning("  Закладки не добавлены → продолжаем без них")
            with_bm = with_blanks

        # ══════════════════════════════════════════════════════════════════
        #  ЭТАП 9: Последняя страница
        # ══════════════════════════════════════════════════════════════════
        final_pdf = out_root / "final_document.pdf"
        if not add_last_page(str(with_bm), str(final_pdf), rl_font):
            log.warning("  Последняя стр. не добавлена → сохраняем as-is")
            shutil.copy2(str(with_bm), str(final_pdf))

        # ── Итоговый отчёт ─────────────────────────────────────────────────
        total  = page_count(str(final_pdf))
        sheets = math.ceil(total / 2)
        log.info("═" * 56)
        log.info("  ✅  ГОТОВО")
        log.info("  Файл    : %s", final_pdf)
        log.info("  Страниц : %d", total)
        log.info("  Листов  : %d  (ceil(%d / 2))", sheets, total)
        log.info("═" * 56)
        return True

    except Exception as exc:
        log.exception("КРИТИЧЕСКАЯ ОШИБКА: %s", exc)
        return False

    finally:
        try:
            shutil.rmtree(tmpdir, ignore_errors=True)
            log.debug("Temp удалён: %s", tmpdir)
        except Exception:
            pass


# ═════════════════════════════════════════════════════════════════════════════
#  ТОЧКА ВХОДА
# ═════════════════════════════════════════════════════════════════════════════

def main() -> None:
    """
    python pdf_assembler.py                    → ./project_source → ./project_source/output
    python pdf_assembler.py ./src              → ./src → ./src/output
    python pdf_assembler.py ./src ./out        → явные пути
    """
    if len(sys.argv) == 3:
        src, out = sys.argv[1], sys.argv[2]
    elif len(sys.argv) == 2:
        src = sys.argv[1]
        out = str(Path(sys.argv[1]) / "output")
    else:
        src = "./project_source"
        out = "./project_source/output"

    sys.exit(0 if main_pipeline(src, out) else 1)


if __name__ == "__main__":
    main()
