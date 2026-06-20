import os
import docx
from docx.shared import Pt
from docx.oxml.ns import qn


def is_dash(val):
    """Проверяет, является ли значение прочерком или пустым"""
    if val is None:
        return True
    val_str = str(val).strip().replace(' ', '')
    return val_str in ('---', '--', '-', '----', '—', '–', '')


def extract_data_from_word(file_path):
    """Парсит исходный .docx и возвращает список записей"""
    try:
        doc = docx.Document(file_path)
        results = []
        current_block = None
        types_order = ['2', '3', '4']

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                # Начало нового блока вещества
                if 'Вещество:' in ''.join(cells):
                    if current_block:
                        results.append(current_block)

                    header = ''.join(cells)
                    code_and_name = header.split(':')[1].strip()
                    if code_and_name.endswith('Вещество'):
                        code_and_name = code_and_name[:-len('Вещество')].strip()
                    parts = code_and_name.split()
                    code = parts[0] if parts else 'Unknown'
                    name = ' '.join(parts[1:]) if len(parts) > 1 else ''

                    current_block = {
                        'Вещество': f'{code} {name}',
                        'Prepr': {},
                        'SZZ': {},
                        'JZ': {}
                    }
                    continue

                # Сбор данных по веществу
                if current_block and len(cells) >= 13:
                    point_num = cells[0]
                    pdk = cells[4]
                    point_type = cells[12]

                    if point_type in types_order:
                        if point_type == '2':
                            type_field = 'Prepr'
                        elif point_type == '3':
                            type_field = 'SZZ'
                        else:
                            type_field = 'JZ'

                        if not current_block.get(type_field):
                            current_block[type_field] = {
                                'номер': point_num,
                                'PDK': pdk
                            }

        if current_block:
            results.append(current_block)

        # Формируем финальный список записей
        final_results = []
        for block in results:
            # Проверка: есть ли хоть одно непустое значение ПДК
            has_valid_pdk = False
            for t in types_order:
                if t == '2':
                    type_field = 'Prepr'
                elif t == '3':
                    type_field = 'SZZ'
                else:
                    type_field = 'JZ'
                entry = block.get(type_field, {})
                pdk_val = entry.get('PDK', '---')
                if not is_dash(pdk_val):
                    has_valid_pdk = True
                    break

            if not has_valid_pdk:
                continue

            for t in types_order:
                if t == '2':
                    type_field = 'Prepr'
                elif t == '3':
                    type_field = 'SZZ'
                else:
                    type_field = 'JZ'

                entry = block.get(type_field, {})

                record = {
                    'Вещество': block['Вещество'],
                    'Номер точки': entry.get('номер', ''),
                    'Предпр.': entry.get('PDK') if t == '2' else '---',
                    'СЗЗ': entry.get('PDK') if t == '3' else '---',
                    'ЖЗ': entry.get('PDK') if t == '4' else '---'
                }
                final_results.append(record)

        return final_results

    except Exception as e:
        print(f"Ошибка при обработке файла {file_path}: {e}")
        return []


def set_cell_font(cell, font_name='Times New Roman', font_size=10):
    """Применяет шрифт ко всем параграфам в ячейке"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # Важно: задаём также восточноазиатский шрифт для корректного отображения
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = docx.oxml.OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = docx.oxml.OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)


def save_to_docx(data, output_path):
    """Сохраняет данные в .docx в виде таблицы без заголовков"""
    doc = docx.Document()
    
    # Убираем лишние отступы у документа (по желанию)
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(10)
        section.bottom_margin = Pt(10)
        section.left_margin = Pt(15)
        section.right_margin = Pt(15)

    if not data:
        doc.add_paragraph("Нет данных для вывода.")
        doc.save(output_path)
        return

    # Создаём таблицу: строк = количество записей, столбцов = 5
    num_cols = 5
    table = doc.add_table(rows=len(data), cols=num_cols)
    table.style = 'Table Grid'  # С рамками

    # Заполняем таблицу
    for i, record in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = str(record.get('Вещество', ''))
        row.cells[1].text = str(record.get('Номер точки', ''))
        row.cells[2].text = str(record.get('Предпр.', ''))
        row.cells[3].text = str(record.get('СЗЗ', ''))
        row.cells[4].text = str(record.get('ЖЗ', ''))

    # Применяем шрифт ко всем ячейкам
    for row in table.rows:
        for cell in row.cells:
            # Выравнивание по центру (опционально)
            for paragraph in cell.paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(cell, 'Times New Roman', 10)

    doc.save(output_path)


def process_folder():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    for filename in os.listdir(current_dir):
        if filename.endswith('.docx') and not filename.startswith('~') and not filename.startswith('output_'):
            file_path = os.path.join(current_dir, filename)
            try:
                print(f"\nОбработка файла: {filename}")
                data = extract_data_from_word(file_path)

                if data:
                    # Формируем имя выходного файла
                    base_name = os.path.splitext(filename)[0]
                    output_filename = f'output_{base_name}.docx'
                    output_path = os.path.join(current_dir, output_filename)

                    save_to_docx(data, output_path)
                    print(f"✅ Данные сохранены в {output_filename}")
                    print(f"   Записей: {len(data)}")
                else:
                    print(f"⚠ Пустые данные для файла {filename}")

            except Exception as e:
                print(f"❌ Ошибка при обработке файла {filename}: {e}")


if __name__ == "__main__":
    try:
        process_folder()
        print("\n✅ Обработка всех файлов завершена")
    except Exception as e:
        print(f"❌ Произошла ошибка: {e}")