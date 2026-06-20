import re
import json
import logging
import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

# ------------------- Настройка логирования -------------------
logging.basicConfig(
    filename='processing.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    encoding='utf-8'
)
logger = logging.getLogger(__name__)

# ------------------- конфигурация -------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Шаблоны для поиска файла вкладов
INPUT_PATTERNS = [
    os.path.join(SCRIPT_DIR, 'Вклады_МР.xlsx'),
    os.path.join(SCRIPT_DIR, 'Вклады_МР.xls'),
    os.path.join(SCRIPT_DIR, 'Вклады МР.xlsx'),
    os.path.join(SCRIPT_DIR, 'Вклады МР.xls'),
    os.path.join(SCRIPT_DIR, '*.xlsx'),
    os.path.join(SCRIPT_DIR, '*.xls'),
    os.path.join(SCRIPT_DIR, '*.html'),
    os.path.join(SCRIPT_DIR, '*.htm'),
]
PDK_JSON = os.path.join(SCRIPT_DIR, 'base_emissons.json')
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, 'Результат_НМУ.docx')

# ------------------- поиск файла вкладов -------------------
def find_input_file():
    import glob
    for pattern in INPUT_PATTERNS:
        if '*' not in pattern and os.path.isfile(pattern):
            return pattern
    for pattern in INPUT_PATTERNS:
        if '*' in pattern:
            files = glob.glob(pattern)
            if files:
                return files[0]
    return None

INPUT_EXCEL = find_input_file()
if INPUT_EXCEL is None:
    msg = 'Не найден файл с вкладами МР (*.xls, *.xlsx, *.html) в папке со скриптом.'
    logger.error(msg)
    print(msg)
    sys.exit(1)
logger.info(f'Используется файл вкладов: {INPUT_EXCEL}')

# ------------------- проверка справочника ПДК -------------------
if not os.path.isfile(PDK_JSON):
    msg = f'Не найден справочник ПДК: {PDK_JSON}'
    logger.error(msg)
    print(msg)
    sys.exit(1)
logger.info(f'Справочник ПДК: {PDK_JSON}')

# ------------------- загрузка справочника ПДК -------------------
try:
    with open(PDK_JSON, 'r', encoding='utf-8') as f:
        substances = json.load(f)
    pdk_dict = {}
    klass_dict = {}
    for s in substances:
        code = s['code']
        pdk = s.get('pdk_mr')
        pdk_dict[code] = pdk if pdk is not None else '-'
        hc = s.get('hazard_class')
        klass_dict[code] = hc if hc is not None else '-'
    logger.info(f'Загружено {len(pdk_dict)} веществ из справочника ПДК.')
except Exception as e:
    logger.error(f'Ошибка загрузки справочника ПДК: {e}')
    raise

# ------------------- чтение исходного файла (Excel / HTML) -------------------
df = None
ext = os.path.splitext(INPUT_EXCEL)[1].lower()

if ext in ('.xls', '.xlsx'):
    # Пробуем calamine
    try:
        df = pd.read_excel(INPUT_EXCEL, sheet_name=0, header=None, dtype=str, engine='calamine')
        logger.info('Прочитано движком calamine.')
    except Exception as e_cal:
        logger.info(f'calamine не сработал: {e_cal}')
    # Если не вышло, пробуем openpyxl для .xlsx
    if df is None and ext == '.xlsx':
        try:
            df = pd.read_excel(INPUT_EXCEL, sheet_name=0, header=None, dtype=str, engine='openpyxl')
            logger.info('Прочитано движком openpyxl.')
        except Exception as e:
            logger.warning(f'openpyxl не сработал: {e}')
    # Запасной вариант – HTML
    if df is None:
        logger.info('Попытка прочитать файл как HTML...')
        try:
            df_list = pd.read_html(INPUT_EXCEL, header=None)
            if df_list:
                df = df_list[0]
                logger.info(f'Прочитано как HTML, таблиц: {len(df_list)} (первая).')
        except Exception as e_html:
            logger.error(f'Чтение HTML не удалось: {e_html}')
    if df is None:
        msg = ('Не удалось прочитать файл ни в одном формате.\n'
               'Сохраните выгрузку из программы как "Документ Excel (XML)" (.xlsx).')
        logger.error(msg)
        print(msg)
        sys.exit(1)
elif ext in ('.html', '.htm'):
    try:
        df = pd.read_html(INPUT_EXCEL, header=None)[0]
        logger.info('Прочитано как HTML.')
    except Exception as e:
        logger.error(f'Ошибка чтения HTML: {e}')
        raise
else:
    msg = f'Неподдерживаемый формат: {ext}. Поддерживаются .xls, .xlsx, .html'
    logger.error(msg)
    print(msg)
    sys.exit(1)

# Приводим столбцы к одному уровню
if isinstance(df.columns, pd.MultiIndex):
    df.columns = range(df.shape[1])
logger.info(f'Файл прочитан, строк: {len(df)}')

# ------------------- поиск границ таблиц -------------------
header_mask = df[0].astype(str).str.contains('Загрязняющее вещество', na=False)
header_indices = df.index[header_mask].tolist()
logger.info(f'Найдено заголовков таблиц: {len(header_indices)}')

if not header_indices:
    msg = 'Не найдены заголовки таблиц с "Загрязняющее вещество"'
    logger.error(msg)
    raise ValueError(msg)

def is_data_row(row):
    first = str(row[0]).strip() if pd.notna(row[0]) else ''
    return bool(re.match(r'^\d{4}\s', first))

all_records = []

for idx, start_hdr in enumerate(header_indices):
    end_hdr = header_indices[idx + 1] if idx + 1 < len(header_indices) else len(df)
    logger.info(f'Обработка блока {idx+1}: строки {start_hdr}-{end_hdr-1}')

    data_start = None
    for i in range(start_hdr, end_hdr):
        cell_val = str(df.iloc[i, 0]).strip()
        if cell_val in ('1', '> 1'):
            data_start = i + 1
            break
    if data_start is None:
        logger.warning(f'Не найдено начало данных в блоке {idx+1}, пропускаем')
        continue

    block_data = df.iloc[data_start:end_hdr].copy()
    block_data.columns = range(block_data.shape[1])

    block_records = 0
    for _, row in block_data.iterrows():
        if not is_data_row(row):
            continue

        raw_za = str(row[0]).strip()
        m = re.match(r'^(\d{4})\s+(.*)', raw_za)
        if not m:
            continue
        code = m.group(1)
        name = m.group(2).strip()

        point = str(row[1]).strip() if pd.notna(row[1]) else ''

        # жилая зона (столбец 6 -> индекс 5)
        raw_zh = str(row[5]).strip() if pd.notna(row[5]) else ''
        nums = re.findall(r'\d+,\d+', raw_zh.replace('.', ','))
        if not nums:
            continue
        nums_float = [float(n.replace(',', '.')) for n in nums]

        if raw_zh.startswith('/') and len(nums_float) >= 2:
            conc_zh = nums_float[1]
        elif '/' in raw_zh and len(nums_float) >= 2:
            conc_zh = nums_float[0]
        else:
            conc_zh = nums_float[0]

        src = str(row[6]).strip() if pd.notna(row[6]) else ''
        perc = str(row[7]).strip().replace(',', '.') if pd.notna(row[7]) else ''

        all_records.append({
            'code': code,
            'name': name,
            'point': point,
            'conc_zh': conc_zh,
            'source': src,
            'perc': perc
        })
        block_records += 1

    logger.info(f'  извлечено записей: {block_records}')

logger.info(f'Всего извлечено записей: {len(all_records)}')

if not all_records:
    msg = 'Не найдено ни одной записи с концентрацией в жилой зоне.'
    logger.error(msg)
    raise ValueError(msg)

# ------------------- группировка: максимум по веществу -------------------
df_rec = pd.DataFrame(all_records)
idx_max = df_rec.groupby('code')['conc_zh'].idxmax()
df_max = df_rec.loc[idx_max].sort_values('code').reset_index(drop=True)
logger.info(f'После группировки веществ: {len(df_max)}')

# ------------------- формирование DOCX (раздельные столбцы Код и Наименование) -------------------
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

doc.add_paragraph(
    'Таблица 1.1 – Результаты расчета концентраций ЗВ для обоснования перечня ЗВ, '
    'подлежащих сокращению в период НМУ',
    style='Normal'
)

num_cols = 11
table = doc.add_table(rows=1, cols=num_cols, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = [
    '№ п/п', 'Код', 'Наименование ЗВ', 'ПДК', 'Класс опасности',
    'Номер контр.\nточки', 'Максимальная приземная концентрация\nв долях ПДК (жилая зона)',
    'Необходимо уменьшение\nвыбросов при НМУ', 'Увеличение приземной\nконцентрации при НМУ (+20%)',
    '№ ИЗАВ', '% вклада'
]
hdr_cells = table.rows[0].cells
for i, text in enumerate(headers):
    hdr_cells[i].text = text
    for paragraph in hdr_cells[i].paragraphs:
        paragraph.style = doc.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.bold = True

for n, (_, row) in enumerate(df_max.iterrows(), start=1):
    cells = table.add_row().cells

    cells[0].text = str(n)
    cells[1].text = row['code']
    cells[2].text = row['name']
    cells[3].text = str(pdk_dict.get(row['code'], '-'))
    cells[4].text = str(klass_dict.get(row['code'], '-'))
    cells[5].text = row['point']

    conc = round(row['conc_zh'], 4)
    cells[6].text = str(conc).replace('.', ',')

    cells[7].text = '-'

    inc = round(conc * 1.2, 4)
    cells[8].text = str(inc).replace('.', ',')

    cells[9].text = row['source']
    cells[10].text = row['perc']

    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = 1   # по центру
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Ширина столбцов (подобрана для 11 колонок)
widths = [0.8, 1.0, 4.5, 1.0, 1.0, 1.3, 3.0, 1.8, 2.2, 1.2, 1.2]
for i, width in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = Cm(width)

doc.save(OUTPUT_DOCX)
logger.info(f'Файл сохранён: {OUTPUT_DOCX}')
print(f'Готово! Результат записан в {OUTPUT_DOCX}')