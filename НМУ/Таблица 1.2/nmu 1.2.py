import json
import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ================= ВСПОМОГАТЕЛЬНЫЕ ФUNКЦИИ =================

def detect_file_type(filepath):
    """Определяет тип файла по внутренним маркерам JSON"""
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
            prog = data.get('options', {}).get('programm', '').upper()
            if 'PDV' in prog:
                return 'PDV'
            elif 'ECO' in prog:
                return 'UPRZA'
            return 'UNKNOWN'
    except Exception:
        return 'UNKNOWN'

def set_cell_formatting(cell, font_name='Times New Roman', font_size=10, bold=False, align_center=True):
    """Применяет шрифт, выравнивание и границы ко всем параграфам в ячейке"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'bottom', 'left', 'right']:
        # ИСПРАВЛЕНО: корректный поиск тегов с префиксом через qn()
        tag_name = qn(f'w:{border_name}')
        for old_border in tcPr.findall(f'.//{tag_name}'):
            tcPr.remove(old_border)
            
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcPr.append(border)

def set_row_height(row, height_cm):
    """Устанавливает минимальную высоту строки таблицы"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old_height in trPr.findall(qn('w:trHeight')):
        trPr.remove(old_height)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567))) # 1 см = 567 twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def clean_zero(val):
    if val is None: return 0.0
    v = float(val)
    return 0.0 if -0.0000001 < v < 0.0000001 else round(v, 7)

def format_val(val):
    if val is None: return "0"
    try: v = float(val)
    except (ValueError, TypeError): return str(val)
    if -0.0000001 < v < 0.0000001: return "0"
    if abs(v) < 1e-6 and v != 0:
        res = f"{v:.2E}".replace("E-0", "E-").replace("E+0", "E+")
    else:
        res = f"{v:.7f}".rstrip('0').rstrip('.')
    return res.replace('.', ',')

def is_dash(val):
    if val is None: return True
    val_str = str(val).strip().replace(' ', '').replace('\n', '')
    return val_str in ('---', '--', '-', '—', '–', '')

def format_efficiency(eff):
    if eff is None or is_dash(eff): return "-"
    try:
        val = float(eff)
        if val == int(val): return str(int(val))
        return f"{val:.1f}".replace('.', ',')
    except (ValueError, TypeError): return "-"

# ================= ПАРСИНГ WORD ОТЧЕТОВ =================

def extract_concentrations_from_word(file_path):
    try:
        doc = Document(file_path)
        results = {}
        current_substance = None
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if 'Вещество:' in ''.join(cells):
                    header = ''.join(cells)
                    code_and_name = header.split(':')[1].strip()
                    if code_and_name.endswith('Вещество'):
                        code_and_name = code_and_name[:-len('Вещество')].strip()
                    parts = code_and_name.split(None, 1)
                    code = parts[0] if parts else 'Unknown'
                    name = parts[1] if len(parts) > 1 else ''
                    current_substance = f"{code} {name}"
                    if current_substance not in results:
                        results[current_substance] = {}
                    continue
                
                if current_substance and len(cells) >= 13:
                    point_num_str = cells[0].strip()
                    pdk = cells[4].strip()
                    point_type = cells[12].strip()
                    
                    if point_num_str.isdigit():
                        point_num = int(point_num_str)
                        
                        # ФИЛЬТР: Убираем ПЗ (тип 2) полностью
                        if point_type == '2':
                            continue
                            
                        suffix = ""
                        if point_type == '3': suffix = " (СЗЗ)"
                        elif point_type == '4': suffix = " (ЖЗ)"
                        
                        display_num = f"{point_num}{suffix}"
                        
                        if not is_dash(pdk):
                            results[current_substance][point_num] = {
                                'display_num': display_num,
                                'pdk': pdk,
                                'type': point_type
                            }
        return results
    except Exception as e:
        print(f"Ошибка при обработке файла {file_path}: {e}")
        return {}

# ================= РАБОТА С JSON УПРЗА =================

def get_emissions_data(upzra_data):
    normal = {}
    nmu_ent = {}
    nmu_meas = {}
    cps_with_measures = set() # Множество веществ, к которым реально применены мероприятия
    
    # 1. Обычные условия
    for ist in upzra_data.get('IST', {}).values():
        if ist.get('IST_VN') == 1:
            for cnt in ist.get('CNT', {}).values():
                cp = cnt.get('CNT_CP')
                if cp: normal[cp] = normal.get(cp, 0) + clean_zero(cnt.get('CNT_GSGRS', 0))
                    
    # 2. Находим источники с мероприятиями (IST_VN=2, FUSE=3 в CVCODE=2)
    active_measure_ist_nns = set()
    for cv in upzra_data.get('ECO_CV', {}).values():
        if cv.get('CODE') == 2:
            for su in cv.get('SELECT_USE', {}).values():
                if su.get('FUSE') == 3 and su.get('IST_VN') == 2:
                    active_measure_ist_nns.add(su.get('IST_NN'))
                    
    # 3. Группируем источники по IST_NN
    ist_by_nn = {}
    for ist in upzra_data.get('IST', {}).values():
        nn = ist.get('IST_NN')
        vn = ist.get('IST_VN')
        if nn not in ist_by_nn: ist_by_nn[nn] = {}
        ist_by_nn[nn][vn] = ist
        
    # 4. Расчет выбросов НМУ
    for nn, versions in ist_by_nn.items():
        ist_data = versions.get(2) or versions.get(1) # Берем версию НМУ, если есть, иначе обычную
        if not ist_data: continue
        
        is_measure_source = (nn in active_measure_ist_nns)
        
        for cnt in ist_data.get('CNT', {}).values():
            cp = cnt.get('CNT_CP')
            if not cp: continue
            gsg = clean_zero(cnt.get('CNT_GSGRS', 0))
            
            # Выбросы по предприятию включают ВСЕ источники в их состоянии НМУ
            nmu_ent[cp] = nmu_ent.get(cp, 0) + gsg
            
            # Выбросы по мероприятию включают ТОЛЬКО источники с мерами
            if is_measure_source:
                nmu_meas[cp] = nmu_meas.get(cp, 0) + gsg
                cps_with_measures.add(cp) # Запоминаем, что у этого вещества есть мероприятия
                
    return normal, nmu_ent, nmu_meas, cps_with_measures

def get_substance_names(upzra_data):
    names = {}
    for cv_data in upzra_data.get('ECO_CV', {}).values():
        for pol_data in cv_data.get('SELECT_POLCNT', {}).values():
            cp = pol_data.get('CODE')
            nm = pol_data.get('NAME')
            if cp and nm: names[cp] = nm.strip()
    return names

def get_sources_for_substance(upzra_data, cp_code):
    sources = set()
    for ist_data in upzra_data.get('IST', {}).values():
        if ist_data.get('IST_VN') == 1:
            for cnt_data in ist_data.get('CNT', {}).values():
                if cnt_data.get('CNT_CP') == cp_code:
                    sources.add(ist_data.get('IST_NN'))
                    break
    return sorted(list(sources))

# ================= ГЕНЕРАЦИЯ ОТЧЕТА WORD =================

def generate_nmu_report(original_upzra, nmu_upzra, concentrations_normal, concentrations_nmu, include_szz, output_file):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    subst_names = get_substance_names(original_upzra)
    normal_emissions, nmu_ent_emissions, nmu_meas_emissions, cps_with_measures = get_emissions_data(nmu_upzra)
    # Для обычных условий берем только normal (первый элемент кортежа)
    temp_norm, _, _, _ = get_emissions_data(original_upzra)
    normal_emissions = temp_norm

    all_cp_codes = set()
    all_cp_codes.update(normal_emissions.keys())
    all_cp_codes.update(nmu_ent_emissions.keys())
    all_cp_codes.update(nmu_meas_emissions.keys())

    for subst_key in list(concentrations_normal.keys()) + list(concentrations_nmu.keys()):
        parts = subst_key.split(None, 1)
        if parts:
            try: all_cp_codes.add(int(parts[0]))
            except ValueError: pass

    sorted_cp_codes = sorted(list(all_cp_codes))
    
    # Создаем ОДНУ сквозную таблицу для всего документа
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    
    headers_emissions = ["Номера источников выброса", "Обычные условия\nВыброс (г/с)", "Режим НМУ\nпо предприятию\nВыброс (г/с)", "Режим НМУ\nпо мероприятию\nВыброс (г/0)"]
    headers_conc = ["Номер контрольной точки", "Концентрация обычные\n(долей ПДК)", "Концентрация НМУ\n(долей ПДК)", "Эффективность (%)"]

    for cp_code in sorted_cp_codes:
        subst_name = subst_names.get(cp_code, str(cp_code))
        
        # ИСПРАВЛЕНИЕ: Форматируем код с ведущими нулями до 4 знаков для поиска в Word (стандарт УПРЗА)
        cp_code_str = f"{cp_code:04d}"
        
        subst_key_normal = next((k for k in concentrations_normal.keys() if k.startswith(f"{cp_code_str} ") or f" {cp_code_str} " in k), None)
        subst_key_nmu = next((k for k in concentrations_nmu.keys() if k.startswith(f"{cp_code_str} ") or f" {cp_code_str} " in k), None)
        
        all_points = set()
        if subst_key_normal: all_points.update(concentrations_normal[subst_key_normal].keys())
        if subst_key_nmu: all_points.update(concentrations_nmu[subst_key_nmu].keys())
        
        filtered_points = []
        for point_num in all_points:
            point_type = None
            if subst_key_normal and point_num in concentrations_normal[subst_key_normal]:
                point_type = concentrations_normal[subst_key_normal][point_num]['type']
            elif subst_key_nmu and point_num in concentrations_nmu[subst_key_nmu]:
                point_type = concentrations_nmu[subst_key_nmu][point_num]['type']
            
            # Фильтр СЗЗ (тип 3) по запросу пользователя (тип 2 уже отброшен в парсере)
            if point_type == '3' and not include_szz: continue
                
            if subst_key_normal and point_num in concentrations_normal[subst_key_normal]:
                filtered_points.append(concentrations_normal[subst_key_normal][point_num])
            elif subst_key_nmu and point_num in concentrations_nmu[subst_key_nmu]:
                filtered_points.append(concentrations_nmu[subst_key_nmu][point_num])
        
        # Сортировка точек по числовому значению
        def sort_key(p):
            num_str = str(p.get('display_num', '')).split()[0]
            return int(num_str) if num_str.isdigit() else 0
        filtered_points.sort(key=sort_key)

        sources = get_sources_for_substance(original_upzra, cp_code)
        sources_str = ', '.join(map(str, sources)) if sources else '-'
        norm_val = normal_emissions.get(cp_code, 0)
        nmu_ent_val = nmu_ent_emissions.get(cp_code, 0)
        
        # ИСПРАВЛЕНИЕ ЛОГИКИ: Если мероприятий для вещества нет, выброс по мероприятию = обычному
        if cp_code in cps_with_measures:
            nmu_meas_val = nmu_meas_emissions.get(cp_code, 0)
        else:
            nmu_meas_val = norm_val 

        eff_ent = ((norm_val - nmu_ent_val) / norm_val * 100) if norm_val > 0 else 0
        eff_meas = ((norm_val - nmu_meas_val) / norm_val * 100) if norm_val > 0 else 0
        
        # --- 1. Заголовок вещества (объединенная ячейка с увеличенной высотой) ---
        row = table.add_row()
        cell = row.cells[0].merge(row.cells[3])
        # УБРАНО: "с учетом фоновых концентраций"
        cell.text = f"Вещество: {cp_code} {subst_name}"
        set_cell_formatting(cell, font_size=11, bold=True, align_center=False)
        set_row_height(row, 1.2) # Визуальное разделение блоков (~1.2 см)

        # --- 2. Шапка выбросов ---
        row = table.add_row()
        for i, h_text in enumerate(headers_emissions):
            cell = row.cells[i]
            cell.text = h_text
            set_cell_formatting(cell, font_size=9, bold=True)

        # --- 3. Данные выбросов ---
        row = table.add_row()
        row.cells[0].text = sources_str
        row.cells[1].text = format_val(norm_val)
        row.cells[2].text = format_val(nmu_ent_val)
        row.cells[3].text = format_val(nmu_meas_val)
        for i in range(4): set_cell_formatting(row.cells[i], font_size=10)

        # --- 4. Эффективность ---
        row = table.add_row()
        row.cells[0].text = "Эффективность мероприятий (%)"
        row.cells[1].text = "-"
        row.cells[2].text = format_efficiency(eff_ent)
        row.cells[3].text = format_efficiency(eff_meas)
        
        set_cell_formatting(row.cells[0], font_size=10, bold=True, align_center=False)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i in range(1, 4): set_cell_formatting(row.cells[i], font_size=10, bold=True)

        # --- 5. Шапка концентраций ---
        row = table.add_row()
        for i, h_text in enumerate(headers_conc):
            cell = row.cells[i]
            cell.text = h_text
            set_cell_formatting(cell, font_size=9, bold=True)

        # --- 6. Точки ---
        for point_data in filtered_points:
            row = table.add_row()
            display_num = point_data.get('display_num', '')
            orig_num = int(display_num.split()[0]) if display_num.split()[0].isdigit() else None
            
            pdk_norm = '-'
            if subst_key_normal and orig_num in concentrations_normal[subst_key_normal]:
                pdk_norm = concentrations_normal[subst_key_normal][orig_num].get('pdk', '-')

            pdk_nmu = '-'
            if subst_key_nmu and orig_num in concentrations_nmu[subst_key_nmu]:
                pdk_nmu = concentrations_nmu[subst_key_nmu][orig_num].get('pdk', '-')

            row.cells[0].text = display_num
            row.cells[1].text = pdk_norm
            row.cells[2].text = pdk_nmu
            
            eff_conc = "-"
            if not is_dash(pdk_norm) and not is_dash(pdk_nmu):
                try:
                    val_norm = float(str(pdk_norm).replace(',', '.'))
                    val_nmu = float(str(pdk_nmu).replace(',', '.'))
                    if val_norm > 0: eff_conc = (val_norm - val_nmu) / val_norm * 100
                except (ValueError, TypeError): pass
            
            row.cells[3].text = format_efficiency(eff_conc)
            for i in range(4): set_cell_formatting(row.cells[i], font_size=10)

    doc.save(output_file)
    print(f"Отчет успешно сохранен: {output_file}")

# ================= ОСНОВНОЙ ПРОЦЕСС =================

def main():
    print("=" * 70)
    print("ГЕНЕРАЦИЯ ОТЧЕТА НМУ (Финальная версия)")
    print("=" * 70)
    
    files_to_process = sys.argv[1:] if len(sys.argv) > 1 else [f for f in os.listdir('.') if f.endswith('.json') or f.endswith('.docx')]

    upzra_normal_file = None
    upzra_nmu_file = None
    report_normal_file = None
    report_nmu_file = None

    for f in files_to_process:
        if f.endswith('.json'):
            ftype = detect_file_type(f)
            if ftype == 'UPRZA':
                if 'нму' in f.lower() or 'nmu' in f.lower(): upzra_nmu_file = f
                else: upzra_normal_file = f
        elif f.endswith('.docx') and not f.startswith('~') and 'отчет' in f.lower():
            if 'нму' in f.lower() or 'nmu' in f.lower(): report_nmu_file = f
            else: report_normal_file = f

    if not upzra_normal_file or not upzra_nmu_file:
        print("Ошибка: Не найдены файлы УПРЗА (исходный и НМУ).")
        return

    print(f"УПРЗА исходный: {upzra_normal_file}")
    print(f"УПРЗА НМУ: {upzra_nmu_file}")
    print(f"Отчет МР (Обычные): {report_normal_file or 'Не найден'}")
    print(f"Отчет МР (НМУ): {report_nmu_file or 'Не найден (концентрации НМУ будут прочерками)'}")

    with open(upzra_normal_file, 'r', encoding='utf-8-sig') as f: original_upzra = json.load(f)
    with open(upzra_nmu_file, 'r', encoding='utf-8-sig') as f: nmu_upzra = json.load(f)

    concentrations_normal = extract_concentrations_from_word(report_normal_file) if report_normal_file else {}
    concentrations_nmu = extract_concentrations_from_word(report_nmu_file) if report_nmu_file else {}

    print("\n" + "=" * 70)
    include_szz_input = input("Учитывать точки СЗЗ (тип 3)? (1 - да, 0 - нет): ").strip()
    include_szz = include_szz_input == '1'
    print(f"Точки СЗЗ: {'включены' if include_szz else 'исключены'}")
    print("=" * 70)

    output_file = "Отчет_НМУ_итоговый.docx"
    print(f"\nГенерация отчета: {output_file} ...")

    generate_nmu_report(original_upzra, nmu_upzra, concentrations_normal, concentrations_nmu, include_szz, output_file)
    print("\nРабота скрипта завершена!")

if __name__ == "__main__":
    main()
