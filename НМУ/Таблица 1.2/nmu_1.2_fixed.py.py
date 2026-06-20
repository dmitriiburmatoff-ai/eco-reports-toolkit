import json
import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ================= ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =================

def detect_file_type(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
            prog = data.get('options', {}).get('programm', '').upper()
            if 'PDV' in prog:
                return 'PDV'
            elif 'ECO' in prog:
                return 'UPRZA'
            return 'UNKNOWN'
    except Exception:
        return 'UNKNOWN'

def set_cell_formatting(cell, font_name='Times New Roman', font_size=10, bold=False, align_center=True, valign='center'):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    if valign == 'center':
        vAlign.set(qn('w:val'), 'center')
    elif valign == 'top':
        vAlign.set(qn('w:val'), 'top')
    elif valign == 'bottom':
        vAlign.set(qn('w:val'), 'bottom')

    for border_name in ['top', 'bottom', 'left', 'right']:
        tag_name = qn(f'w:{border_name}')
        for old_border in tcPr.findall(f'.//{tag_name}'):
            tcPr.remove(old_border)
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcPr.append(border)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old_height in trPr.findall(qn('w:trHeight')):
        trPr.remove(old_height)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

# ИСПРАВЛЕНИЕ: порог обнуления снижен до 1e-12
def clean_zero(val):
    if val is None: return 0.0
    v = float(val)
    return 0.0 if -1e-12 < v < 1e-12 else round(v, 12)  # сохраняем до 12 знаков

def format_val(val):
    if val is None: return "0"
    try: v = float(val)
    except (ValueError, TypeError): return str(val)
    if -1e-12 < v < 1e-12: return "0"
    if abs(v) < 1e-6 and v != 0:
        res = f"{v:.2E}".replace("E-0", "E-").replace("E+0", "E+")
    else:
        res = f"{v:.7f}".rstrip('0').rstrip('.')
    return res.replace('.', ',')

def is_dash(val):
    if val is None: return True
    val_str = str(val).strip().replace(' ', '').replace('\n', '')
    return val_str in ('---', '--', '-', '—', '–', '')

def format_efficiency(eff):
    if eff is None or is_dash(eff): return "-"
    try:
        val = float(eff)
        if val == int(val): return str(int(val))
        return f"{val:.1f}".replace('.', ',')
    except (ValueError, TypeError): return "-"

def normalize_key(key):
    return ''.join(key.split())

# ================= ПАРСИНГ WORD ОТЧЕТОВ =================

def extract_concentrations_from_word(file_path):
    try:
        doc = Document(file_path)
        results = {}
        current_substance = None
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = ''.join(cells)
                if 'Вещество:' in row_text:
                    parts = row_text.split('Вещество:')
                    if len(parts) > 1:
                        code_and_name = parts[1].strip()
                        if code_and_name.endswith('\\'):
                            code_and_name = code_and_name[:-1].strip()
                        if code_and_name.endswith('Вещество'):
                            code_and_name = code_and_name[:-len('Вещество')].strip()
                        sub_parts = code_and_name.split(None, 1)
                        code = sub_parts[0] if sub_parts else 'Unknown'
                        name = sub_parts[1] if len(sub_parts) > 1 else ''
                        code = code.strip()
                        current_substance = f"{code} {name}" if name else code
                        if current_substance not in results:
                            results[current_substance] = {}
                        continue
                
                if current_substance and len(cells) >= 13:
                    point_num_str = cells[0].strip()
                    pdk = cells[4].strip()
                    point_type = cells[12].strip()
                    
                    if point_num_str.isdigit():
                        point_num = int(point_num_str)
                        if point_type == '2':
                            continue
                        suffix = ""
                        if point_type == '3':
                            suffix = " (СЗЗ)"
                        elif point_type == '4':
                            suffix = " (ЖЗ)"
                        display_num = f"{point_num}{suffix}"
                        results[current_substance][point_num] = {
                            'display_num': display_num,
                            'pdk': pdk,
                            'type': point_type
                        }
        return results
    except Exception as e:
        print(f"Ошибка при обработке файла {file_path}: {e}")
        return {}

# ================= РАБОТА С JSON УПРЗА =================

def get_emissions_data(upzra_data):
    normal = {}
    nmu_ent = {}
    nmu_meas = {}
    cps_with_measures = set()
    
    for ist in upzra_data.get('IST', {}).values():
        if ist.get('IST_VN') == 1:
            for cnt in ist.get('CNT', {}).values():
                cp = cnt.get('CNT_CP')
                if cp: normal[cp] = normal.get(cp, 0) + clean_zero(cnt.get('CNT_GSGRS', 0))
                    
    active_measure_ist_nns = set()
    for cv in upzra_data.get('ECO_CV', {}).values():
        if cv.get('CODE') == 2:
            for su in cv.get('SELECT_USE', {}).values():
                if su.get('FUSE') == 3 and su.get('IST_VN') == 2:
                    active_measure_ist_nns.add(su.get('IST_NN'))
                    
    ist_by_nn = {}
    for ist in upzra_data.get('IST', {}).values():
        nn = ist.get('IST_NN')
        vn = ist.get('IST_VN')
        if nn not in ist_by_nn: ist_by_nn[nn] = {}
        ist_by_nn[nn][vn] = ist
        
    for nn, versions in ist_by_nn.items():
        ist_data = versions.get(2) or versions.get(1)
        if not ist_data: continue
        
        is_measure_source = (nn in active_measure_ist_nns)
        
        for cnt in ist_data.get('CNT', {}).values():
            cp = cnt.get('CNT_CP')
            if not cp: continue
            gsg = clean_zero(cnt.get('CNT_GSGRS', 0))
            
            nmu_ent[cp] = nmu_ent.get(cp, 0) + gsg
            
            if is_measure_source:
                nmu_meas[cp] = nmu_meas.get(cp, 0) + gsg
                cps_with_measures.add(cp)
                
    return normal, nmu_ent, nmu_meas, cps_with_measures

def get_substance_names(upzra_data):
    names = {}
    for cv_data in upzra_data.get('ECO_CV', {}).values():
        for pol_data in cv_data.get('SELECT_POLCNT', {}).values():
            cp = pol_data.get('CODE')
            nm = pol_data.get('NAME')
            if cp and nm: names[cp] = nm.strip()
    return names

def get_sources_for_substance(upzra_data, cp_code):
    sources = set()
    for ist_data in upzra_data.get('IST', {}).values():
        if ist_data.get('IST_VN') == 1:
            for cnt_data in ist_data.get('CNT', {}).values():
                if cnt_data.get('CNT_CP') == cp_code:
                    sources.add(ist_data.get('IST_NN'))
                    break
    return sorted(list(sources))

# ================= ГЕНЕРАЦИЯ ОТЧЕТА WORD =================

def generate_nmu_report(original_upzra, nmu_upzra, concentrations_normal, concentrations_nmu, include_szz, output_file):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    subst_names = get_substance_names(original_upzra)
    normal_emissions, nmu_ent_emissions, nmu_meas_emissions, cps_with_measures = get_emissions_data(nmu_upzra)
    temp_norm, _, _, _ = get_emissions_data(original_upzra)
    normal_emissions = temp_norm

    all_cp_codes = set()
    all_cp_codes.update(normal_emissions.keys())
    all_cp_codes.update(nmu_ent_emissions.keys())
    all_cp_codes.update(nmu_meas_emissions.keys())

    for subst_key in list(concentrations_normal.keys()) + list(concentrations_nmu.keys()):
        parts = subst_key.split(None, 1)
        if parts:
            try: all_cp_codes.add(int(parts[0]))
            except ValueError: pass

    sorted_cp_codes = sorted(list(all_cp_codes))
    
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    
    # ИСПРАВЛЕНИЕ: заменил "г/0" на "г/с"
    headers_emissions = ["Номера источников выброса", "Обычные условия\nВыброс (г/с)", "Режим НМУ\nпо предприятию\nВыброс (г/с)", "Режим НМУ\nпо мероприятию\nВыброс (г/с)"]
    headers_conc = ["Номер контрольной точки", "Концентрация обычные\n(долей ПДК)", "Концентрация НМУ\n(долей ПДК)", "Эффективность (%)"]

    print(f"Найдено веществ: {len(sorted_cp_codes)}")

    for cp_code in sorted_cp_codes:
        subst_name = subst_names.get(cp_code, str(cp_code))
        cp_code_str = f"{cp_code:04d}"
        
        # Поиск ключа в словарях концентраций
        norm_target = cp_code_str
        subst_key_normal = None
        subst_key_nmu = None
        for key in concentrations_normal.keys():
            if normalize_key(key).startswith(norm_target):
                subst_key_normal = key
                break
        for key in concentrations_nmu.keys():
            if normalize_key(key).startswith(norm_target):
                subst_key_nmu = key
                break
        
        if subst_key_normal is None:
            norm_target_nozero = str(int(cp_code_str))
            for key in concentrations_normal.keys():
                if normalize_key(key).startswith(norm_target_nozero):
                    subst_key_normal = key
                    break
        if subst_key_nmu is None:
            norm_target_nozero = str(int(cp_code_str))
            for key in concentrations_nmu.keys():
                if normalize_key(key).startswith(norm_target_nozero):
                    subst_key_nmu = key
                    break

        all_points = set()
        if subst_key_normal:
            all_points.update(concentrations_normal[subst_key_normal].keys())
        if subst_key_nmu:
            all_points.update(concentrations_nmu[subst_key_nmu].keys())
        
        if not all_points:
            print(f"Для вещества {cp_code_str} не найдено точек в отчётах.")

        filtered_points = []
        for point_num in all_points:
            point_type = None
            if subst_key_normal and point_num in concentrations_normal[subst_key_normal]:
                point_type = concentrations_normal[subst_key_normal][point_num]['type']
            elif subst_key_nmu and point_num in concentrations_nmu[subst_key_nmu]:
                point_type = concentrations_nmu[subst_key_nmu][point_num]['type']
            
            if point_type == '3' and not include_szz:
                continue
                
            if subst_key_normal and point_num in concentrations_normal[subst_key_normal]:
                filtered_points.append(concentrations_normal[subst_key_normal][point_num])
            elif subst_key_nmu and point_num in concentrations_nmu[subst_key_nmu]:
                filtered_points.append(concentrations_nmu[subst_key_nmu][point_num])
        
        def sort_key(p):
            num_str = str(p.get('display_num', '')).split()[0]
            return int(num_str) if num_str.isdigit() else 0
        filtered_points.sort(key=sort_key)

        sources = get_sources_for_substance(original_upzra, cp_code)
        sources_str = ', '.join(f"{s:04d}" for s in sources) if sources else '-'
        norm_val = normal_emissions.get(cp_code, 0)
        nmu_ent_val = nmu_ent_emissions.get(cp_code, 0)
        
        if cp_code in cps_with_measures:
            nmu_meas_val = nmu_meas_emissions.get(cp_code, 0)
        else:
            nmu_meas_val = norm_val 

        eff_ent = ((norm_val - nmu_ent_val) / norm_val * 100) if norm_val > 0 else 0
        eff_meas = ((norm_val - nmu_meas_val) / norm_val * 100) if norm_val > 0 else 0
        
        # --- Заголовок вещества ---
        row = table.add_row()
        cell = row.cells[0].merge(row.cells[3])
        cell.text = f"Вещество: {cp_code_str} {subst_name}"
        set_cell_formatting(cell, font_size=11, bold=True, align_center=False, valign='center')
        set_row_height(row, 0.6)

        # --- Шапка выбросов ---
        row = table.add_row()
        for i, h_text in enumerate(headers_emissions):
            cell = row.cells[i]
            cell.text = h_text
            set_cell_formatting(cell, font_size=9, bold=False)

        # --- Данные выбросов ---
        row = table.add_row()
        row.cells[0].text = sources_str
        row.cells[1].text = format_val(norm_val)
        row.cells[2].text = format_val(nmu_ent_val)
        row.cells[3].text = format_val(nmu_meas_val)
        for i in range(4):
            set_cell_formatting(row.cells[i], font_size=10, bold=False)

        # --- Эффективность выбросов ---
        row = table.add_row()
        row.cells[0].text = "Эффективность мероприятий (%)"
        row.cells[1].text = "-"
        row.cells[2].text = format_efficiency(eff_ent)
        row.cells[3].text = format_efficiency(eff_meas)
        
        set_cell_formatting(row.cells[0], font_size=10, bold=False, align_center=False)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i in range(1, 4):
            set_cell_formatting(row.cells[i], font_size=10, bold=False)

        # --- Шапка концентраций ---
        row = table.add_row()
        for i, h_text in enumerate(headers_conc):
            cell = row.cells[i]
            cell.text = h_text
            set_cell_formatting(cell, font_size=9, bold=False)

        # --- Точки ---
        if not filtered_points:
            row = table.add_row()
            row.cells[0].text = "Нет данных"
            row.cells[1].text = "-"
            row.cells[2].text = "-"
            row.cells[3].text = "-"
            for i in range(4):
                set_cell_formatting(row.cells[i], font_size=10, bold=False)
        else:
            for point_data in filtered_points:
                row = table.add_row()
                display_num = point_data.get('display_num', '')
                orig_num = int(display_num.split()[0]) if display_num.split()[0].isdigit() else None
                
                pdk_norm = '-'
                if subst_key_normal and orig_num in concentrations_normal[subst_key_normal]:
                    pdk_norm = concentrations_normal[subst_key_normal][orig_num].get('pdk', '-')

                if eff_ent >= 99.999:  # эффективность по предприятию 100%
                    pdk_nmu = "0"
                    eff_conc = 100.0
                else:
                    pdk_nmu = '-'
                    if subst_key_nmu and orig_num in concentrations_nmu[subst_key_nmu]:
                        pdk_nmu = concentrations_nmu[subst_key_nmu][orig_num].get('pdk', '-')
                    
                    eff_conc = "-"
                    if not is_dash(pdk_norm) and not is_dash(pdk_nmu):
                        try:
                            val_norm = float(str(pdk_norm).replace(',', '.'))
                            val_nmu = float(str(pdk_nmu).replace(',', '.'))
                            if val_norm > 0:
                                eff_conc = (val_norm - val_nmu) / val_norm * 100
                        except (ValueError, TypeError):
                            pass

                row.cells[0].text = display_num
                row.cells[1].text = pdk_norm
                row.cells[2].text = pdk_nmu
                row.cells[3].text = format_efficiency(eff_conc)
                for i in range(4):
                    set_cell_formatting(row.cells[i], font_size=10, bold=False)

    doc.save(output_file)
    print(f"Отчет успешно сохранен: {output_file}")

# ================= ОСНОВНОЙ ПРОЦЕСС =================

def main():
    print("=" * 70)
    print("ГЕНЕРАЦИЯ ОТЧЕТА НМУ (финальная версия)")
    print("=" * 70)
    
    files_to_process = sys.argv[1:] if len(sys.argv) > 1 else [f for f in os.listdir('.') if f.endswith('.json') or f.endswith('.docx')]

    upzra_normal_file = None
    upzra_nmu_file = None
    report_normal_file = None
    report_nmu_file = None

    for f in files_to_process:
        if 'итоговый' in f.lower():
            continue
        if f.endswith('.json'):
            ftype = detect_file_type(f)
            if ftype == 'UPRZA':
                if 'нму' in f.lower() or 'nmu' in f.lower():
                    upzra_nmu_file = f
                else:
                    upzra_normal_file = f
        elif f.endswith('.docx') and not f.startswith('~'):
            if 'отчет мр' in f.lower() and 'нму' not in f.lower():
                report_normal_file = f
            elif 'отчет мр' in f.lower() and 'нму' in f.lower():
                report_nmu_file = f

    if not report_normal_file:
        for f in files_to_process:
            if f.endswith('.docx') and 'отчет' in f.lower() and 'нму' not in f.lower() and 'итоговый' not in f.lower():
                report_normal_file = f
                break
    if not report_nmu_file:
        for f in files_to_process:
            if f.endswith('.docx') and 'отчет' in f.lower() and 'нму' in f.lower() and 'итоговый' not in f.lower():
                report_nmu_file = f
                break

    if not upzra_normal_file or not upzra_nmu_file:
        print("Ошибка: Не найдены файлы УПРЗА (исходный и НМУ).")
        return

    print(f"УПРЗА исходный: {upzra_normal_file}")
    print(f"УПРЗА НМУ: {upzra_nmu_file}")
    print(f"Отчет МР (Обычные): {report_normal_file or 'Не найден'}")
    print(f"Отчет МР (НМУ): {report_nmu_file or 'Не найден (концентрации НМУ будут прочерками)'}")

    with open(upzra_normal_file, 'r', encoding='utf-8-sig') as f:
        original_upzra = json.load(f)
    with open(upzra_nmu_file, 'r', encoding='utf-8-sig') as f:
        nmu_upzra = json.load(f)

    concentrations_normal = extract_concentrations_from_word(report_normal_file) if report_normal_file else {}
    concentrations_nmu = extract_concentrations_from_word(report_nmu_file) if report_nmu_file else {}

    print(f"\nИз обычного отчёта извлечено веществ: {len(concentrations_normal)}")
    print(f"Из отчёта НМУ извлечено веществ: {len(concentrations_nmu)}")

    print("\n" + "=" * 70)
    include_szz_input = input("Учитывать точки СЗЗ (тип 3)? (1 - да, 0 - нет): ").strip()
    include_szz = include_szz_input == '1'
    print(f"Точки СЗЗ: {'включены' if include_szz else 'исключены'}")
    print("=" * 70)

    output_file = "Отчет_НМУ_итоговый.docx"
    print(f"\nГенерация отчета: {output_file} ...")

    generate_nmu_report(original_upzra, nmu_upzra, concentrations_normal, concentrations_nmu, include_szz, output_file)
    print("\nРабота скрипта завершена!")

if __name__ == "__main__":
    main()