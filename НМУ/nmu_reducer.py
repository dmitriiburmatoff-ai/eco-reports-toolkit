import json
import os
import sys
import logging
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Настройка логирования
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('nmu_processing.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# ================= ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =================

def detect_file_type(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
        prog = data.get('options', {}).get('programm', '').upper()
        if 'PDV' in prog:
            return 'PDV'
        elif 'ECO' in prog:
            return 'UPRZA'
        return 'UNKNOWN'
    except Exception as e:
        logger.error(f"Ошибка определения типа файла {filepath}: {e}")
        return 'UNKNOWN'

def set_cell_vertical_alignment(cell, align):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)
    tcPr.append(valign)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcPr.append(border)

def clean_zero(val):
    if val is None:
        return 0.0
    v = float(val)
    # Порог 1e-11 сохраняет значения бенз(а)пирена (порядка 1e-8)
    if -1e-11 < v < 1e-11:
        return 0.0
    return round(v, 11)

def format_val(val):
    if val is None:
        return "0"
    try:
        v = float(val)
    except (ValueError, TypeError):
        return str(val)
    
    if -1e-11 < v < 1e-11:
        return "0"

    if abs(v) < 1e-6 and v != 0:
        res = f"{v:.2E}".replace("E-0", "E-").replace("E+0", "E+")
    else:
        res = f"{v:.7f}".rstrip('0').rstrip('.')
        
    return res.replace('.', ',')

def find_vyd_by_id(data, vyd_id):
    for key, val in data.get('VYD', {}).items():
        if val.get('VYD_ID') == vyd_id:
            return key, val
    return None, None

def find_ist_by_nn_vn(data, ist_nn, ist_vn):
    for key, val in data.get('IST', {}).items():
        if val.get('IST_NN') == ist_nn and val.get('IST_VN') == ist_vn:
            return key, val
    return None, None

def build_subst_map(pdv_data, upzra_data):
    """Собирает карту кодов веществ и их названий"""
    subst_map = {}
    for ss_data in pdv_data.get('ShortSubst', {}).values():
        cp = ss_data.get('CODE')
        nm = ss_data.get('NAME')
        if cp and nm: 
            subst_map[cp] = nm
    
    for ist_data in pdv_data.get('IST', {}).values():
        for cnt_data in ist_data.get('CNT', {}).values():
            cp = cnt_data.get('CNT_CP')
            nm = cnt_data.get('CNT_NAME')
            if cp and nm and cp not in subst_map:
                subst_map[cp] = nm
        
    for cv_data in upzra_data.get('ECO_CV', {}).values():
        for pol_data in cv_data.get('SELECT_POLCNT', {}).values():
            cp = pol_data.get('CODE')
            nm = pol_data.get('NAME')
            if cp and nm and cp not in subst_map:
                subst_map[cp] = nm
        
    for ist_data in upzra_data.get('IST', {}).values():
        for cnt_data in ist_data.get('CNT', {}).values():
            cp = cnt_data.get('CNT_CP')
            nm = cnt_data.get('CNT_NAME')
            if cp and nm and cp not in subst_map:
                subst_map[cp] = nm

    logger.info(f"Всего собрано веществ: {len(subst_map)}")
    return subst_map

# ================= ГЕНЕРАЦИЯ ОТЧЁТА WORD =================

def generate_report(pdv_data, initial_ist_state, all_ists, affected_ists, subst_map):
    logger.info("Начало генерации отчета Word")
    print("\nГенерация отчёта об изменениях...")
    doc = Document()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Отчёт об изменении выбросов загрязняющих веществ (НМУ)")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = "Times New Roman"

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    headers = ["№ п/п", "№ ИЗАВ", " ", "Вещество (Код - Наименование)", "Выброс г/с (до)", "Выброс г/с (после)"]
    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = header_cells[i]
        cell.text = text
        set_cell_vertical_alignment(cell, 'center')
        set_cell_borders(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.font.bold = True
            if i in [0, 1, 4, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    seq_num = 1
    current_row_idx = 1
    sorted_ists = sorted(list(all_ists), key=lambda x: (x[0], x[1]))

    for ist_nn, ist_vn in sorted_ists:
        if ist_vn != 1:
            continue
            
        ist_key, ist_data = find_ist_by_nn_vn(pdv_data, ist_nn, ist_vn)
        if not ist_data:
            continue
            
        substances = []
        ist_cnt = ist_data.get('CNT', {})
        unique_ist_id = (ist_nn, ist_vn)
        
        for cp_key, cp_data in ist_cnt.items():
            cp_code = cp_data.get('CNT_CP')
            before_gsg = initial_ist_state.get(unique_ist_id, {}).get(cp_code, {}).get('gsg', 0)
            after_gsg = cp_data.get('CNT_GSGRS', 0) 
            
            name = subst_map.get(cp_code, str(cp_code))
            
            substances.append({
                'code': cp_code,
                'name': name,
                'before': before_gsg,
                'after': after_gsg
            })
            
        num_subst = len(substances)
        if num_subst == 0:
            continue
            
        for _ in range(num_subst):
            table.add_row()
         
        if num_subst > 1:
            table.cell(current_row_idx, 0).merge(table.cell(current_row_idx + num_subst - 1, 0))
            table.cell(current_row_idx, 1).merge(table.cell(current_row_idx + num_subst - 1, 1))
            table.cell(current_row_idx, 2).merge(table.cell(current_row_idx + num_subst - 1, 2))
            
        for i, subst in enumerate(substances):
            row = table.rows[current_row_idx + i]
            cells = row.cells
            
            if i == 0:
                cells[0].text = str(seq_num)
                cells[1].text = f"{ist_nn:04d}"
                cells[2].text = " "
                
            cells[3].text = f"{subst['code']} - {subst['name']}"
            cells[4].text = format_val(subst['before'])
            cells[5].text = format_val(subst['after'])
            
            for col_idx, cell in enumerate(cells):
                set_cell_vertical_alignment(cell, 'center')
                set_cell_borders(cell)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(10)
                        r.font.bold = False
                    if col_idx in [0, 1, 4, 5]:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
        current_row_idx += num_subst
        seq_num += 1

    report_filename = "Отчёт_НМУ_изменения.docx"
    doc.save(report_filename)
    logger.info(f"Отчет сохранен: {report_filename}")
    print(f"Отчёт успешно сохранён: {report_filename}")

# ================= СИНХРОНИЗАЦИЯ УПРЗА =================

def update_upzra(upzra_data, original_pdv, reduced_pdv, affected_ists):
    logger.info("Начало синхронизации УПРЗА")
    print("\nСинхронизация данных УПРЗА...")
    
    original_ists = {val.get('IST_NN'): val for val in original_pdv.get('IST', {}).values() if val.get('IST_NN') is not None}
    reduced_ists = {val.get('IST_NN'): val for val in reduced_pdv.get('IST', {}).values() if val.get('IST_NN') is not None}
    ist_section = upzra_data.get('IST', {})

    # 1. Обновляем основные источники (IST_VN=1) исходными значениями из original_pdv
    logger.info("Шаг 1: Обновление основных источников (VN=1) исходными значениями")
    for ist_nn, src in original_ists.items():
        for ist in ist_section.values():
            if ist.get('IST_NN') == ist_nn and ist.get('IST_VN') == 1:
                if 'CNT' in src:
                    src_cnt = {cdata.get('CNT_CP'): cdata for cdata in src['CNT'].values()}
                    for cnt_data in ist.get('CNT', {}).values():
                        cp_code = cnt_data.get('CNT_CP')
                        if cp_code in src_cnt:
                            new_cnt = src_cnt[cp_code]
                            cnt_data['CNT_GSGRS'] = clean_zero(new_cnt.get('CNT_GSGRS', 0))
                            cnt_data['CNT_TG'] = clean_zero(new_cnt.get('CNT_TG', 0))
                            if 'CNT_GRS' in cnt_data and 'CNT_GRS' in new_cnt:
                                cnt_data['CNT_GRS'] = clean_zero(new_cnt.get('CNT_GRS', 0))
                break

    # 2. Собираем информацию о существующих VN=2
    logger.info("Шаг 2: Анализ существующих вариантов НМУ (VN=2)")
    existing_vn2 = {}
    for key, ist in ist_section.items():
        if ist.get('IST_VN') == 2:
            ist_nn = ist.get('IST_NN')
            existing_vn2[ist_nn] = key

    # 3. Создаем/обновляем НМУ варианты (VN=2) ТОЛЬКО для измененных источников
    logger.info("Шаг 3: Создание/обновление НМУ вариантов (VN=2) только для измененных источников")
    logger.debug(f"Измененные источники: {affected_ists}")

    for ist_nn, ist_vn in affected_ists:
        if ist_vn != 1:
            continue
        
        src = reduced_ists.get(ist_nn)
        if not src:
            continue

        # Ищем основной источник в УПРЗА
        base_ist = next((ist for ist in ist_section.values() if ist.get('IST_NN') == ist_nn and ist.get('IST_VN') == 1), None)
        if not base_ist:
            logger.warning(f"Основной источник IST_NN={ist_nn} (VN=1) не найден в УПРЗА. Пропускаем.")
            continue

        # Проверяем, есть ли уже VN=2
        if ist_nn in existing_vn2:
            copy_key = existing_vn2[ist_nn]
            copy_ist = ist_section[copy_key]
            logger.debug(f"Обновляем существующий НМУ вариант: IST_NN={ist_nn}")
        else:
            # Создаем новый с VN=2
            copy_ist = deepcopy(base_ist)
            copy_ist['IST_VN'] = 2
            base_name = str(base_ist.get('IST_NAME', ''))
            if not base_name.endswith(' (НМУ)'):
                copy_ist['IST_NAME'] = base_name + ' (НМУ)'
            copy_ist['DT'] = '1899-12-30'
            max_rec = max([int(k[3:]) for k in ist_section.keys() if k.startswith('REC')], default=0)
            copy_key = f"REC{max_rec+1}"
            ist_section[copy_key] = copy_ist
            existing_vn2[ist_nn] = copy_key
            logger.debug(f"Создан новый НМУ вариант: IST_NN={ist_nn}, ключ={copy_key}")
            
        # Обновляем выбросы в копии (берем из reduced_pdv, т.е. уже с учетом снижений)
        if 'CNT' in src:
            src_cnt = {cdata.get('CNT_CP'): cdata for cdata in src['CNT'].values()}
            for cnt_data in copy_ist.get('CNT', {}).values():
                cp_code = cnt_data.get('CNT_CP')
                if cp_code in src_cnt:
                    new_cnt = src_cnt[cp_code]
                    cnt_data['CNT_GSGRS'] = clean_zero(new_cnt.get('CNT_GSGRS', 0))
                    cnt_data['CNT_TG'] = clean_zero(new_cnt.get('CNT_TG', 0))
                    if 'CNT_GRS' in cnt_data and 'CNT_GRS' in new_cnt:
                        cnt_data['CNT_GRS'] = clean_zero(new_cnt.get('CNT_GRS', 0))

    logger.info(f"Всего создано/обновлено НМУ вариантов (VN=2): {len(existing_vn2)}")

    # 4. Обновляем SELECT_USE для CVCODE=2
    logger.info("Шаг 4: Обновление SELECT_USE для CVCODE=2")
    eco_cv = upzra_data.get('ECO_CV', {})
    nmu_cv = next((cv for cv in eco_cv.values() if cv.get('CODE') == 2), None)

    if nmu_cv is None:
        logger.warning("Вариант CVCODE=2 не найден, создаем новый")
        first_cv = next(iter(eco_cv.values())) if eco_cv else None
        if first_cv:
            nmu_cv = deepcopy(first_cv)
            nmu_cv['CODE'] = 2
            nmu_cv['NAME'] = "Мероприятия при НМУ"
            max_cv_rec = max([int(k[3:]) for k in eco_cv.keys() if k.startswith('REC')], default=0)
            eco_cv[f"REC{max_cv_rec+1}"] = nmu_cv
        else:
            logger.error("Нет вариантов расчёта в УПРЗА")
            return False

    select_use = nmu_cv.get('SELECT_USE', {})
    select_use.clear()
    rec_counter = 0

    # Сначала добавляем все VN=2 (только измененные) с FUSE=3
    logger.debug("Добавляем источники VN=2 с FUSE=3")
    for ist in ist_section.values():
        if ist.get('IST_VN') == 2:
            ist_nn = ist.get('IST_NN')
            base_ist = next((bi for bi in ist_section.values() if bi.get('IST_NN') == ist_nn and bi.get('IST_VN') == 1), None)
            xcode = base_ist.get('XCODE', '') if base_ist else ist.get('XCODE', '')
            plo_nn = ist.get('PLO_NN', 1)
            plo_name = upzra_data.get('PLO', {}).get(f"REC{plo_nn-1}", {}).get('NAME', 'Промплощадка')
            ent_name = upzra_data.get('PRE', {}).get('NAME', '')
            
            select_use[f"REC{rec_counter}"] = {
                "FUSE": 3, "PRIORITET": 0, "IST_NN": ist_nn, "IST_VN": 2, "XCODE": xcode,
                "IST_NAME": ist.get('IST_NAME', ''), "IST_TP": ist.get('IST_TP', 1),
                "PLOCODE": plo_nn, "PLONAME": plo_name, "CECHCODE": ist.get('CH_NN', 1),
                "CECHNAME": " ", "CVCODE": 2, "CVNAME": "Мероприятия при НМУ",
                "VER_CODE": 1, "VER_NAME": "Существующее положение", "ENT_CODE": 1, "ENT_NAME": ent_name
            }
            rec_counter += 1

    # Затем добавляем все VN=1 с правильными флагами FUSE
    logger.debug("Добавляем источники VN=1 с правильными флагами FUSE")
    for ist in ist_section.values():
        if ist.get('IST_VN') == 1:
            ist_nn = ist.get('IST_NN')
            plo_nn = ist.get('PLO_NN', 1)
            plo_name = upzra_data.get('PLO', {}).get(f"REC{plo_nn-1}", {}).get('NAME', 'Промплощадка')
            ent_name = upzra_data.get('PRE', {}).get('NAME', '')
            
            # Если источник был изменен (есть в affected_ists), то FUSE=-1 (не использовать базовые значения)
            # Если источник не был изменен, то FUSE=1 (использовать базовые значения)
            if (ist_nn, 1) in affected_ists:
                fuse_val = -1
            else:
                fuse_val = 1
            
            select_use[f"REC{rec_counter}"] = {
                "FUSE": fuse_val, "PRIORITET": 0, "IST_NN": ist_nn, "IST_VN": 1, "XCODE": ist.get('XCODE', ''),
                "IST_NAME": ist.get('IST_NAME', ''), "IST_TP": ist.get('IST_TP', 1),
                "PLOCODE": plo_nn, "PLONAME": plo_name, "CECHCODE": ist.get('CH_NN', 1),
                "CECHNAME": " ", "CVCODE": 2, "CVNAME": "Мероприятия при НМУ",
                "VER_CODE": 1, "VER_NAME": "Существующее положение", "ENT_CODE": 1, "ENT_NAME": ent_name
            }
            rec_counter += 1
            
    nmu_cv['SELECT_USE'] = select_use
    logger.info(f"Всего записей в SELECT_USE для НМУ: {len(select_use)}")

    return True

# ================= ОСНОВНОЙ ПРОЦЕСС =================

def main():
    logger.info("="*70)
    logger.info("НАЧАЛО РАБОТЫ СКРИПТА")
    logger.info("="*70)
    files_to_process = sys.argv[1:] if len(sys.argv) > 1 else [f for f in os.listdir('.') if f.endswith('.json')]
    pdv_file = None
    upzra_file = None

    for f in files_to_process:
        ftype = detect_file_type(f)
        if ftype == 'PDV':
            pdv_file = f
        elif ftype == 'UPRZA':
            upzra_file = f

    if not pdv_file or not upzra_file:
        logger.error("Не удалось найти файлы ПДВ и УПРЗА")
        print("Ошибка: Не удалось найти файлы ПДВ (маркер 'PDV') и УПРЗА (маркер 'ECO').")
        print("Использование: python nmu_reducer.py <файл1.json> <файл2.json>")
        return

    print(f"Найден файл ПДВ: {pdv_file}")
    print(f"Найден файл УПРЗА: {upzra_file}")

    with open(pdv_file, 'r', encoding='utf-8-sig') as f:
        original_pdv = json.load(f)
    with open(upzra_file, 'r', encoding='utf-8-sig') as f:
        upzra_data = json.load(f)

    working_pdv = deepcopy(original_pdv)
    subst_map = build_subst_map(original_pdv, upzra_data)

    initial_ist_state = {}
    all_ists = set()

    for ist_data in working_pdv.get('IST', {}).values():
        ist_nn = ist_data.get('IST_NN')
        ist_vn = ist_data.get('IST_VN')
        unique_ist_id = (ist_nn, ist_vn)
        
        if ist_vn == 1:
            all_ists.add((ist_nn, ist_vn))
        
        initial_ist_state[unique_ist_id] = {}
        for cp_data in ist_data.get('CNT', {}).values():
            cp_code = cp_data.get('CNT_CP')
            initial_ist_state[unique_ist_id][cp_code] = {'gsg': cp_data.get('CNT_GSGRS', 0)}

    vyb_vyd = working_pdv.get('VYB_VYD', {})
    affected_ists = set()
    vyd_links = {}

    for link_data in vyb_vyd.values():
        vyd_id = link_data.get('VYD_ID')
        if vyd_id not in vyd_links:
            vyd_links[vyd_id] = []
        vyd_links[vyd_id].append({
            'IST_NN': link_data.get('IST_NN'),
            'IST_VN': link_data.get('IST_VN'),
            'K_VYD_IST': link_data.get('K_VYD_IST', 100)
        })

    print("\n" + "="*70)
    print("НАЧАЛО ОБРАБОТКИ СВЯЗЕЙ ИВ -> ИЗАВ")
    print("="*70)

    for vyd_id, links in vyd_links.items():
        vyd_key, vyd_data = find_vyd_by_id(working_pdv, vyd_id)
        if not vyd_data:
            continue

        vyd_code = vyd_data.get('CODE', '?')
        vyd_vn = vyd_data.get('VN', '?')
        vyd_name = vyd_data.get('NAME', 'Без названия')
        
        print(f"\n[ИВ] №{vyd_code} (вар. {vyd_vn}) | Название: '{vyd_name}'")
        print("Связанные ИЗАВ:")
        
        for link in links:
            ist_nn = link['IST_NN']
            ist_vn = link['IST_VN']
            k_vyd = link['K_VYD_IST']
            _, ist_data = find_ist_by_nn_vn(working_pdv, ist_nn, ist_vn)
            ist_name = ist_data.get('IST_NAME', 'Не найдено') if ist_data else 'Не найдено'
            print(f"  -> №{ist_nn:04d} (вар. {ist_vn}): '{ist_name}' (доля: {k_vyd}%)")

        while True:
            user_input = input("\nВведите процент снижения для этого ИВ (0 для пропуска, Enter = 0): ").strip()
            if not user_input:
                percent = 0.0
                break
            try:
                percent = float(user_input)
                if percent < 0:
                    print("Процент не может быть отрицательным.")
                    continue
                if percent > 100:
                    print("Процент не может быть больше 100. Будет использовано 100.")
                    percent = 100.0
                break
            except ValueError:
                print("Неверный формат числа.")

        if percent <= 0:
            print("  >>> Пропущено без изменений.\n" + "-"*70)
            continue

        for link in links:
            affected_ists.add((link['IST_NN'], link['IST_VN']))

        reduction_factor = percent / 100.0
        print(f"  >>> Применяется снижение на {percent}%...")

        decreases = {}
        for cp_data in vyd_data.get('CNTV', {}).values():
            cp_code = cp_data.get('CNT_CP')
            decreases[cp_code] = {
                'gsg': cp_data.get('CNT_GSGRS', 0) * reduction_factor,
                'tg': cp_data.get('CNT_TG', 0) * reduction_factor,
                'grs': cp_data.get('CNT_GRS', 0) * reduction_factor if 'CNT_GRS' in cp_data else 0.0
            }

        for cp_data in vyd_data.get('CNTV', {}).values():
            cp_code = cp_data.get('CNT_CP')
            dec = decreases[cp_code]
            cp_data['CNT_GSGRS'] = clean_zero(cp_data.get('CNT_GSGRS', 0) - dec['gsg'])
            cp_data['CNT_TG'] = clean_zero(cp_data.get('CNT_TG', 0) - dec['tg'])
            if 'CNT_GRS' in cp_data:
                cp_data['CNT_GRS'] = clean_zero(cp_data.get('CNT_GRS', 0) - dec['grs'])

        for link in links:
            ist_nn, ist_vn, k_factor = link['IST_NN'], link['IST_VN'], link['K_VYD_IST'] / 100.0
            ist_key, ist_data = find_ist_by_nn_vn(working_pdv, ist_nn, ist_vn)
            if not ist_data:
                continue

            for cp_data in vyd_data.get('CNTV', {}).values():
                cp_code = cp_data.get('CNT_CP')
                dec = decreases[cp_code]
                for iv_data in ist_data.get('CNT', {}).values():
                    if iv_data.get('CNT_CP') == cp_code:
                        iv_data['CNT_GSGRS'] = clean_zero(iv_data.get('CNT_GSGRS', 0) - (dec['gsg'] * k_factor))
                        iv_data['CNT_TG'] = clean_zero(iv_data.get('CNT_TG', 0) - (dec['tg'] * k_factor))
                        if 'CNT_GRS' in iv_data:
                            iv_data['CNT_GRS'] = clean_zero(iv_data.get('CNT_GRS', 0) - (dec['grs'] * k_factor))
                        break
        print("  >>> Успешно пересчитано.")
        print("-" * 70)

    # Генерация отчета
    generate_report(working_pdv, initial_ist_state, all_ists, affected_ists, subst_map)

    # Синхронизация УПРЗА
    success = update_upzra(upzra_data, original_pdv, working_pdv, affected_ists)

    if success:
        dir_name = os.path.dirname(upzra_file) or "."
        base_name = os.path.basename(upzra_file)
        name_without_ext = os.path.splitext(base_name)[0]
        output_upzra = os.path.join(dir_name, f"{name_without_ext}_НМУ.json")
        
        with open(output_upzra, 'w', encoding='utf-8-sig') as f:
            json.dump(upzra_data, f, ensure_ascii=False, indent=2)
        print(f"\nУПРЗА успешно обновлён и сохранён: {output_upzra}")

    print("\nРабота скрипта завершена!")
    print("Подробный лог сохранен в файле: nmu_processing.log")

if __name__ == "__main__":
    main()
