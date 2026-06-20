#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для автоматического анализа отчётов программы "Эколог-Шум" (версия 2.6.x)
Формирует сводный отчёт с группировкой источников и выводами о превышениях ПДУ.
"""
import os
import re
import sys
from collections import defaultdict
from datetime import datetime
from typing import List, Dict, Optional, Tuple, Any
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

# ----------------------------------------------------------------------
# 1. КОНСТАНТЫ
# ----------------------------------------------------------------------
PDU = {
    'day': {
        '31.5': 90, '63': 75, '125': 66, '250': 59, '500': 54,
        '1000': 50, '2000': 47, '4000': 45, '8000': 44,
        'La.экв': 55, 'La.макс': 70
    },
    'night': {
        '31.5': 83, '63': 67, '125': 57, '250': 49, '500': 44,
        '1000': 40, '2000': 37, '4000': 35, '8000': 33,
        'La.экв': 45, 'La.макс': 60
    }
}

OCTAVE_BANDS = ['31.5', '63', '125', '250', '500', '1000', '2000', '4000', '8000']
EXTRA_PARAMS = ['La.экв', 'La.макс']
ALL_PARAMS = OCTAVE_BANDS + EXTRA_PARAMS
DAY_KEYWORDS = ['день', 'ден', 'дн', 'day']
NIGHT_KEYWORDS = ['ночь', 'ноч', 'нч', 'night']
REPORT_KEYWORDS = ['Эколог-Шум', 'Copyright ©', 'ИНТЕГРАЛ']
HEADER_RESULTS_3_1 = '3.1. Результаты в расчетных точках'

# ----------------------------------------------------------------------
# 2. ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ----------------------------------------------------------------------
def clean_number(text: str) -> Optional[float]:
    if not text: return None
    text = str(text).strip().replace('*', '').replace(',', '.').replace('−', '-').replace('—', '-')
    try: return float(text)
    except ValueError: return None

def normalize_name(name: str) -> Tuple[str, str]:
    name = name.strip().rstrip('.')
    patterns = [
        (r'(.*?)\s*[\.\(]\s*(день|ден|дн|day)\s*[\)\.]?$', 'day'),
        (r'(.*?)\s*[\.\(]\s*(ночь|ноч|нч|night)\s*[\)\.]?$', 'night')
    ]
    for pattern, suffix in patterns:
        match = re.search(pattern, name, re.IGNORECASE)
        if match: return match.group(1).strip(), suffix
    return name, ''

def determine_period_from_filename(filename: str) -> Optional[str]:
    base = os.path.splitext(os.path.basename(filename))[0].lower()
    if any(kw in base for kw in DAY_KEYWORDS): return 'day'
    if any(kw in base for kw in NIGHT_KEYWORDS): return 'night'
    return None

def determine_period_from_header(doc: Document) -> Optional[str]:
    for para in doc.paragraphs:
        text = para.text.lower()
        if 'в дневное время с 7:00 до 23:00' in text: return 'day'
        if 'в ночное время с 23:00 до 7:00' in text: return 'night'
    return None

def is_ecolog_report(doc: Document) -> bool:
    full_text = '\n'.join([p.text for p in doc.paragraphs[:50]])
    return any(kw.lower() in full_text.lower() for kw in REPORT_KEYWORDS)

def find_docx_files(directory: str) -> List[str]:
    return [os.path.join(directory, f) for f in os.listdir(directory)
            if f.lower().endswith('.docx') and not f.startswith('~')]

def set_document_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

def set_cell_font(cell, text, bold=False, color=None, font_size=10):
    cell.text = ''
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color: run.font.color.rgb = color

def get_next_table(paragraph) -> Optional[Any]:
    next_elem = paragraph._element.getnext()
    while next_elem is not None:
        if next_elem.tag.endswith('tbl'):
            for table in paragraph.part.document.tables:
                if table._element is next_elem: return table
        elif next_elem.tag.endswith('p'):
            p_text = ''.join([t.text for t in next_elem.xpath('.//w:t') if t.text])
            if p_text.strip(): break
        next_elem = next_elem.getnext()
    return None

# ----------------------------------------------------------------------
# 3. РАСПОЗНАВАНИЕ И ПАРСИНГ ИСТОЧНИКОВ (Исправлено)
# ----------------------------------------------------------------------
def is_source_table(table) -> bool:
    if len(table.rows) < 2: return False
    header_text = ' '.join(cell.text for row in table.rows[:3] for cell in row.cells)
    has_freq = any(band in header_text for band in ['31.5', '63', '125'])
    has_obj = 'объект' in header_text.lower()
    return has_freq and has_obj

def parse_source_table(table, is_nonconstant: bool) -> List[Dict]:
    sources = []
    header_cells = []
    for r in range(min(3, len(table.rows))):
        header_cells.extend([(c, cell.text.strip()) for c, cell in enumerate(table.rows[r].cells)])

    def find_col(keywords):
        for col, txt in header_cells:
            if any(kw.lower() in txt.lower() for kw in keywords): return col
        return None

    idx_n = find_col(['N', '№', 'n'])
    if idx_n is None: idx_n = 0
    idx_name = find_col(['объект', 'наименование'])
    idx_leq = find_col(['la.экв', 'la экв', 'la.eq'])
    idx_lmax = find_col(['la.макс', 'la макс', 'la.max'])
    idx_octaves = {band: find_col([band]) for band in OCTAVE_BANDS if find_col([band]) is not None}

    if idx_name is None or len(idx_octaves) < 5: return sources

    data_start = 0
    for i, row in enumerate(table.rows):
        if idx_n < len(row.cells):
            val = row.cells[idx_n].text.strip()
            if val.replace('.', '').replace('-', '').isdigit():
                data_start = i; break

    for row in table.rows[data_start:]:
        cells = row.cells
        if idx_n >= len(cells): continue
        n_val = cells[idx_n].text.strip()
        if not n_val or not n_val.replace('.', '').replace('-', '').isdigit(): continue

        name_val = cells[idx_name].text.strip() if idx_name < len(cells) else ''
        levels = {}
        for band, idx in idx_octaves.items():
            if idx < len(cells):
                val = clean_number(cells[idx].text)
                if val is not None: levels[band] = val
        if idx_leq is not None and idx_leq < len(cells):
            val = clean_number(cells[idx_leq].text)
            if val is not None: levels['La.экв'] = val
        if is_nonconstant and idx_lmax is not None and idx_lmax < len(cells):
            val = clean_number(cells[idx_lmax].text)
            if val is not None: levels['La.макс'] = val

        sources.append({'N': n_val, 'name': name_val, 'levels': levels, 'is_nonconstant': is_nonconstant})
    return sources

def extract_sources(doc: Document) -> List[Dict]:
    sources = []
    
    def get_next_table_elem(elem):
        """Находит следующую таблицу после элемента, останавливаясь на текстовых параграфах."""
        next_elem = elem.getnext()
        while next_elem is not None:
            if next_elem.tag.endswith('tbl'):
                return Table(next_elem, doc._body)
            elif next_elem.tag.endswith('p'):
                p_text = ''.join([t.text for t in next_elem.xpath('.//w:t') if t.text])
                if p_text.strip():
                    return None
            next_elem = next_elem.getnext()
        return None

    def process_section(header_text, is_nonconstant):
        for para in doc.paragraphs:
            if header_text in para.text.replace('\n', ' '):
                table = get_next_table(para)
                while table:
                    if is_source_table(table):
                        sources.extend(parse_source_table(table, is_nonconstant))
                    table = get_next_table_elem(table._element)
                break

    process_section('1.2. Источники постоянного шума', False)
    process_section('1.3. Источники непостоянного шума', True)
    return sources

# ----------------------------------------------------------------------
# 4. РАСЧЁТНЫЕ ТОЧКИ
# ----------------------------------------------------------------------
def parse_points_table(table, zone: str) -> List[Dict]:
    points = []
    col_indices = {}
    for row in table.rows[:2]:
        for i, cell in enumerate(row.cells):
            text = cell.text.strip()
            for band in ALL_PARAMS:
                if band in text: col_indices[band] = i
    if not col_indices: return points

    data_start = 0
    for i, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if cells and cells[0].isdigit(): data_start = i; break

    for row in table.rows[data_start:]:
        cells = [c.text.strip() for c in row.cells]
        if not cells or not cells[0].isdigit(): continue
        levels = {band: clean_number(cells[idx]) for band, idx in col_indices.items() 
                  if idx < len(cells) and clean_number(cells[idx]) is not None}
        points.append({'N': cells[0], 'zone': zone, 'levels': levels})
    return points

def extract_points(doc: Document) -> List[Dict]:
    points = []
    in_section = False
    current_zone = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if HEADER_RESULTS_3_1 in text: in_section = True; continue
        if in_section and text.startswith('3.2.'): break
        if in_section and 'Точки типа:' in text:
            if 'производственной зоны' in text: current_zone = 'ПЗ'
            elif 'санитарно-защитной зоны' in text: current_zone = 'СЗЗ'
            elif 'жилой зоны' in text: current_zone = 'ЖЗ'
            table = get_next_table(para)
            if table and current_zone: points.extend(parse_points_table(table, current_zone))
    return points

# ----------------------------------------------------------------------
# 5. ГРУППИРОВКА ИСТОЧНИКОВ
# ----------------------------------------------------------------------
def source_key(src: Dict) -> str:
    base_name, _ = normalize_name(src['name'])
    base_name = base_name.rstrip(' .')
    levels = src['levels']
    parts = [base_name]
    # La.макс ИСКЛЮЧЁН из ключа сравнения
    for param in OCTAVE_BANDS + ['La.экв']:
        val = levels.get(param)
        parts.append(str(val) if val is not None else 'None')
    return '|'.join(parts)

def get_params_without_lmax(levels: Dict) -> Dict:
    return {k: v for k, v in levels.items() if k != 'La.макс'}

def group_sources(sources: List[Dict]) -> List[Dict]:
    groups = defaultdict(list)
    for src in sources: groups[source_key(src)].append(src)

    grouped = []
    for key, src_list in groups.items():
        first = src_list[0]
        base_name, _ = normalize_name(first['name'])
        base_name = base_name.rstrip(' .')
        
        params_list = [get_params_without_lmax(s['levels']) for s in src_list]
        all_same = all(p == params_list[0] for p in params_list)
        
        if not all_same:
            for s in src_list:
                period = s.get('period', '')
                final_name = s['name'].rstrip('.')
                if period == 'day': final_name += ' (день)'
                elif period == 'night': final_name += ' (ночь)'
                
                try: num = int(s['N'])
                except: continue
                
                grouped.append({
                    'name': final_name, 'numbers': f"{num:03d}", 'count': 1, 
                    'levels': s['levels'], 'first_num': num
                })
        else:
            numbers = []
            for s in src_list:
                try: numbers.append(int(s['N']))
                except (ValueError, TypeError): pass
            if not numbers: continue

            unique_numbers = sorted(set(numbers))
            num_str = ', '.join([f"{n:03d}" for n in unique_numbers]) if len(unique_numbers) <= 5 else f"{unique_numbers[0]:03d}-{unique_numbers[-1]:03d}"
            count = len(unique_numbers)
            
            lmax_vals = [s['levels'].get('La.макс') for s in src_list if s['levels'].get('La.макс') is not None]
            levels = first['levels'].copy()
            levels['La.макс'] = max(lmax_vals) if lmax_vals else None
            
            periods = {s['period'] for s in src_list if 'period' in s}
            final_name = base_name
            if len(periods) > 1: final_name += ' (день/ночь)'
            elif len(periods) == 1:
                period = list(periods)[0]
                if period == 'day': final_name += ' (день)'
                elif period == 'night': final_name += ' (ночь)'
            
            grouped.append({
                'name': final_name, 'numbers': num_str, 'count': count, 
                'levels': levels, 'first_num': unique_numbers[0]
            })
    
    grouped.sort(key=lambda x: x['first_num'])
    return grouped

# ----------------------------------------------------------------------
# 6. АНАЛИЗ И ДОЛИ ПДУ
# ----------------------------------------------------------------------
def calculate_pdu_share(value: Optional[float], pdu: Optional[float]) -> Optional[float]:
    if value is None or pdu is None: return None
    return 10 ** ((value - pdu) / 10)

def analyze_points(points: List[Dict], period: str) -> List[Dict]:
    pdu = PDU[period]
    analyzed = []
    for pt in points:
        new_pt = pt.copy()
        new_pt['exceedances'] = {}
        new_pt['shares'] = {}
        for param in ALL_PARAMS:
            val = pt['levels'].get(param)
            pdu_val = pdu.get(param)
            if val is not None and pdu_val is not None:
                new_pt['exceedances'][param] = val >= pdu_val
                new_pt['shares'][param] = calculate_pdu_share(val, pdu_val)
            else:
                new_pt['exceedances'][param] = False
                new_pt['shares'][param] = None
        analyzed.append(new_pt)
    return analyzed

# ----------------------------------------------------------------------
# 7. ВЫВОДЫ
# ----------------------------------------------------------------------
def get_top_two_by_zone(points: List[Dict], zone: str) -> List[Tuple]:
    candidates = []
    for pt in points:
        if pt['zone'] != zone: continue
        for param in ALL_PARAMS:
            val = pt['levels'].get(param)
            share = pt['shares'].get(param)
            if val is not None and share is not None:
                candidates.append((pt, param, val, share))
    if not candidates: return []
    candidates.sort(key=lambda x: (-x[3], OCTAVE_BANDS.index(x[1]) if x[1] in OCTAVE_BANDS else 999))
    return candidates[:2]

def format_share(share):
    if share is None: return "–"
    if 0 < share < 0.005: return f"{share:.4f}".replace('.', ',')
    return f"{share:.2f}".replace('.', ',')

def format_point_output(pt: Dict, param: str, value: float, share) -> str:
    rt_num = pt['N']
    if param == 'La.экв': p_str, unit = 'по эквивалентному уровню звука', 'дБА'
    elif param == 'La.макс': p_str, unit = 'по максимальному уровню звука', 'дБА'
    else: p_str, unit = f'в октавной полосе {param} Гц', 'дБ'
    return f"- расчетная точка РТ {rt_num} {p_str} – {value:.1f}".replace('.',',') + f" {unit} ({format_share(share)} ПДУ);"

def generate_conclusions(points_day: List[Dict], points_night: List[Dict]) -> Dict[str, str]:
    all_shares = [s for pt in points_day + points_night for s in pt['shares'].values() if s is not None]
    max_share = max(all_shares) if all_shares else 0.0

    if max_share < 0.1:
        main_text = """По результатам выполненных расчётов уровней шума на границе промышленной площадки в наиболее неблагоприятных условиях распространения звука установлено:

во всех расчётных точках в дневное и ночное время уровни звукового давления и эквивалентного уровня звука составляют менее 0,1 ПДУ, то есть доля ПДУ <0,1.

В соответствии с пунктом 72 СанПиН 2.1.3684-21, критерием отнесения объекта к источникам воздействия на среду обитания и здоровье человека является превышение 0,1 ПДУ. Поскольку расчётные значения шума не превышают 0,1 ПДУ, рассматриваемая площадка не является источником физического воздействия (шума). Организация санитарно-защитной зоны по шумовому фактору не требуется."""
    elif max_share <= 1.0:
        main_text = """По результатам выполненных расчётов уровней шума на границе промышленной площадки в наиболее неблагоприятных условиях распространения звука установлено:

в дневное время в расчётных точках уровень шума составляет более 0,1 ПДУ, но менее 1 ПДУ;

в ночное время уровень шума также превышает 0,1 ПДУ, но не превышает 1 ПДУ;

превышения предельно допустимого уровня (1 ПДУ) не зафиксировано.

Согласно пункту 72 СанПиН 2.1.3684-21, объект признаётся источником воздействия на среду обитания и здоровье человека по физическому фактору (шум), если уровни шума за пределами промышленной площадки превышают 0,1 ПДУ. Так как расчётные значения превышают 0,1 ПДУ, площадка является объектом негативного воздействия. Превышения ПДУ не выявлено, однако в соответствии с требованиями санитарного законодательства требуется организация санитарно-защитной зоны расчётной величины."""
    else:
        main_text = """По результатам выполненных расчётов уровней шума на границе промышленной площадки в наиболее неблагоприятных условиях распространения звука установлено:

в дневное время в расчётных точках уровень шума превышает ПДУ (доля ПДУ >1);

в ночное время также зафиксировано превышение ПДУ.

В соответствии с пунктом 72 СанПиН 2.1.3684-21, превышение 0,1 ПДУ является основанием для отнесения объекта к источникам воздействия. Кроме того, выявленное превышение ПДУ (более 1 ПДУ) свидетельствует о несоблюдении гигиенических нормативов, установленных СанПиН 1.2.3685-21. Площадка признаётся объектом негативного воздействия на окружающую среду по физическому фактору (шум)."""

    lines_zones = ["Наибольшие значения долей ПДУ на границах санитарно-защитной и жилой зон:"]
    for period_txt, pts in [("В дневное время:", points_day), ("В ночное время:", points_night)]:
        lines_zones.append(period_txt)
        for zone in ['СЗЗ', 'ЖЗ']:
            zone_name = "Санитарно-защитная зона:" if zone == 'СЗЗ' else "Жилая зона:"
            lines_zones.append(zone_name)
            top = get_top_two_by_zone(pts, zone)
            if top:
                for pt, param, val, share in top:
                    lines_zones.append(format_point_output(pt, param, val, share))
            else:
                lines_zones.append("- данные отсутствуют;")
        lines_zones.append("")

    return {'main': main_text, 'zones': '\n'.join(lines_zones)}

# ----------------------------------------------------------------------
# 8. ГЕНЕРАЦИЯ ОТЧЁТА
# ----------------------------------------------------------------------
def create_table_from_data(doc, headers, data, title=None, font_size=10):
    if title:
        p = doc.add_paragraph(title); p.style = 'Heading 2'
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
    for i, hdr in enumerate(headers): set_cell_font(table.rows[0].cells[i], hdr, bold=True, font_size=font_size)
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if isinstance(value, tuple):
                text, exceed = value
                set_cell_font(row_cells[i], text, bold=exceed, color=RGBColor(255,0,0) if exceed else None, font_size=font_size)
            else:
                set_cell_font(row_cells[i], str(value) if value is not None else '', font_size=font_size)
    doc.add_paragraph()

def build_sources_table_data(grouped_sources):
    headers = ['N пп', 'Наименование', '№ ИШ', 'Кол-во, шт.'] + ALL_PARAMS
    data = []
    for i, s in enumerate(grouped_sources, 1):
        row = [str(i), s['name'], s['numbers'], str(s['count'])]
        for p in ALL_PARAMS:
            val = s['levels'].get(p)
            row.append(f"{val:.1f}".replace('.',',') if val is not None else '-')
        data.append(row)
    return headers, data

def build_points_table_data(points):
    headers = ['N пп', 'РТ №', 'Зона'] + ALL_PARAMS
    points_sorted = sorted(points, key=lambda p: int(p['N']))
    data = []
    for i, p in enumerate(points_sorted, 1):
        row = [str(i), p['N'], p['zone']]
        for par in ALL_PARAMS:
            val = p['levels'].get(par)
            exc = p['exceedances'].get(par, False)
            row.append((f"{val:.1f}".replace('.',','), exc) if val is not None else ('-', False))
        data.append(row)
    return headers, data

def generate_report(grouped_sources, points_day, points_night, conclusions):
    doc = Document()
    set_document_margins(doc)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    
    doc.add_heading('Результаты анализа шумового воздействия', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y")}')
    
    doc.add_heading('1. Источники шума', level=1)
    if grouped_sources:
        h, d = build_sources_table_data(grouped_sources)
        create_table_from_data(doc, h, d, font_size=10)
    else: doc.add_paragraph("Источники шума не найдены.")
        
    doc.add_heading('2. Результаты расчета в контрольных точках', level=1)
    if points_day:
        doc.add_heading('Дневное время (7:00–23:00)', level=2)
        create_table_from_data(doc, *build_points_table_data(points_day), font_size=10)
    if points_night:
        doc.add_heading('Ночное время (23:00–7:00)', level=2)
        create_table_from_data(doc, *build_points_table_data(points_night), font_size=10)
        
    doc.add_heading('3. Выводы', level=1)
    doc.add_paragraph(conclusions['main'])
    doc.add_paragraph(conclusions['zones'])
    
    out_name = f"Анализ_шума_от_{datetime.now().strftime('%d_%m_%Y')}.docx"
    doc.save(out_name)
    print(f"Отчёт сохранён в файл: {out_name}")

# ----------------------------------------------------------------------
# 9. ОСНОВНАЯ ЛОГИКА
# ----------------------------------------------------------------------
def main():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    docx_files = find_docx_files(current_dir)
    if not docx_files: print("Не найдено файлов .docx."); sys.exit(1)
    
    reports = [f for f in docx_files if is_ecolog_report(Document(f))]
    if not reports: print("Не найдено отчётов Эколог-Шум."); sys.exit(1)
    
    file_periods = {}
    for fpath in reports:
        fn = os.path.basename(fpath)
        p_fn = determine_period_from_filename(fn)
        p_hdr = determine_period_from_header(Document(fpath))
        file_periods[fpath] = p_fn or p_hdr or 'day'

    all_sources, points_day, points_night = [], [], []
    for fpath, period in file_periods.items():
        print(f"Обработка: {os.path.basename(fpath)} ({'день' if period=='day' else 'ночь'})")
        doc = Document(fpath)
        
        doc_sources = extract_sources(doc)
        for src in doc_sources:
            src['period'] = period
        all_sources.extend(doc_sources)
        
        pts = extract_points(doc)
        analyzed = analyze_points(pts, period)
        if period == 'day': points_day.extend(analyzed)
        else: points_night.extend(analyzed)

    generate_report(group_sources(all_sources), points_day, points_night, 
                   generate_conclusions(points_day, points_night))

if __name__ == '__main__':
    main()
